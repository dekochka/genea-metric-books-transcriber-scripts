#!/usr/bin/env python3
"""
Genealogical Record Transcription Tool

Transcribes handwritten birth, death, and marriage records from Google Drive images 
using Vertex AI Gemini 2.5 Pro and creates formatted Google Docs.

Key Features:
- Processes images from Google Drive folders with flexible naming patterns
- Uses external prompt files for different record types (see prompts/ folder)
- Comprehensive logging and error recovery
- Batch processing with configurable start/count parameters

Configuration:
- Requires a YAML config file (see config/config.yaml.example)
- Pass config file as mandatory argument: python transcribe.py config/your_config.yaml
- Configure image_start_number and image_count for selective processing
- Enable retry_mode to reprocess failed images

For detailed setup instructions, prerequisites, and troubleshooting, see README.md
"""

import io
import os
import sys
import argparse
import logging
import base64
import json
import traceback
import yaml
from abc import ABC, abstractmethod
from datetime import datetime
from typing import Any
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google import genai
from google.genai import types
from google.genai.errors import ClientError, ServerError

# Try to import python-docx for Word output
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    print("WARNING: python-docx not installed. Word output will not be available.")
    print("Install with: pip install python-docx>=0.8.11")

# ------------------------- CONFIGURATION LOADING -------------------------

def load_config(config_path: str) -> dict:
    """
    Load configuration from YAML file with mode detection and validation.
    
    Args:
        config_path: Path to the YAML configuration file
        
    Returns:
        Dictionary containing configuration values (normalized)
        
    Raises:
        FileNotFoundError: If config file doesn't exist
        yaml.YAMLError: If config file is invalid YAML
        ValueError: If configuration validation fails
    """
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Configuration file not found: {config_path}")
    
    with open(config_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)
    
    if config is None:
        raise ValueError(f"Configuration file is empty: {config_path}")
    
    # Detect mode
    mode = detect_mode(config)
    
    # Normalize config (convert legacy to nested structure)
    config = normalize_config(config, mode)
    
    # Validate configuration
    is_valid, errors = validate_config(config, mode)
    if not is_valid:
        raise ValueError(f"Configuration validation failed: {', '.join(errors)}")
    
    return config


def detect_mode(config: dict) -> str:
    """
    Detect mode from configuration.
    
    Detection order:
    1. Explicit mode field → use it
    2. Legacy config (project_id/drive_folder_id at root) → googlecloud
    3. local section present → local
    4. Default → googlecloud (backward compatibility)
    
    Args:
        config: Configuration dictionary
        
    Returns:
        Mode string: 'local' or 'googlecloud'
        
    Raises:
        ValueError: If explicit mode value is invalid
    """
    # Check for explicit mode field
    if 'mode' in config:
        mode = config['mode'].lower()
        if mode not in ['local', 'googlecloud']:
            raise ValueError(f"Invalid mode value: '{mode}'. Must be 'local' or 'googlecloud'")
        return mode
    
    # Legacy detection: check for Google Cloud specific fields at root level
    if 'project_id' in config or 'drive_folder_id' in config:
        return 'googlecloud'
    
    # Check for local section
    if 'local' in config:
        return 'local'
    
    # Default for backward compatibility
    return 'googlecloud'


def normalize_config(config: dict, mode: str) -> dict:
    """
    Normalize configuration to internal format.
    
    Converts legacy flat structure to nested structure for googlecloud mode.
    Ensures all required fields are present with defaults.
    
    Args:
        config: Configuration dictionary (may be legacy format)
        mode: Detected mode ('local' or 'googlecloud')
        
    Returns:
        Normalized configuration dictionary
    """
    normalized = config.copy()
    
    if mode == 'googlecloud':
        # If googlecloud section doesn't exist, create it from legacy flat structure
        if 'googlecloud' not in normalized:
            normalized['googlecloud'] = {}
            gc_config = normalized['googlecloud']
            
            # Move Google Cloud specific fields to nested structure
            if 'project_id' in normalized:
                gc_config['project_id'] = normalized.pop('project_id')
            if 'drive_folder_id' in normalized:
                gc_config['drive_folder_id'] = normalized.pop('drive_folder_id')
            if 'region' in normalized:
                gc_config['region'] = normalized.pop('region', 'global')
            if 'ocr_model_id' in normalized:
                gc_config['ocr_model_id'] = normalized.pop('ocr_model_id')
            if 'adc_file' in normalized:
                gc_config['adc_file'] = normalized.pop('adc_file')
            if 'document_name' in normalized:
                gc_config['document_name'] = normalized.pop('document_name')
            if 'title_page_filename' in normalized:
                gc_config['title_page_filename'] = normalized.pop('title_page_filename')
        else:
            # Ensure googlecloud section has defaults
            gc_config = normalized['googlecloud']
            if 'region' not in gc_config:
                gc_config['region'] = 'global'
    
    # Set defaults for shared fields
    if 'retry_mode' not in normalized:
        normalized['retry_mode'] = False
    if 'retry_image_list' not in normalized:
        normalized['retry_image_list'] = []
    
    return normalized


def validate_config(config: dict, mode: str) -> tuple[bool, list[str]]:
    """
    Validate configuration based on mode.
    
    Args:
        config: Configuration dictionary
        mode: Mode string ('local' or 'googlecloud')
        
    Returns:
        Tuple of (is_valid, error_messages)
    """
    errors = []
    
    # Shared required fields
    # Support both prompt_file (legacy) and prompt_template (wizard mode)
    if 'prompt_file' not in config and 'prompt_template' not in config:
        errors.append("Either prompt_file or prompt_template must be specified")
    elif 'prompt_template' in config:
        # Validate template exists
        template_name = config['prompt_template']
        templates_dir = os.path.join(os.path.dirname(__file__), "prompts", "templates")
        template_path = os.path.join(templates_dir, f"{template_name}.md")
        if not os.path.exists(template_path):
            errors.append(f"Prompt template not found: {template_path}")
        # Note: Context validation will be added in Phase 3
        # For now, context is optional (wizard may not have collected it yet)
    
    if 'archive_index' not in config:
        errors.append("archive_index is required")
    
    # Shared optional fields with validation
    if 'image_start_number' in config:
        if not isinstance(config['image_start_number'], int) or config['image_start_number'] < 1:
            errors.append("image_start_number must be a positive integer")
    
    if 'image_count' in config:
        if not isinstance(config['image_count'], int) or config['image_count'] < 1:
            errors.append("image_count must be a positive integer")
    
    # Mode-specific validation
    if mode == 'local':
        if 'local' not in config:
            errors.append("local mode requires 'local' configuration section")
        else:
            local_config = config['local']
            
            # API key validation (can be in config or env var)
            api_key = local_config.get('api_key') or os.getenv('GEMINI_API_KEY')
            if not api_key:
                errors.append("local mode requires api_key in config or GEMINI_API_KEY environment variable")
            elif isinstance(api_key, str) and len(api_key) < 10:
                errors.append("api_key appears to be invalid (too short)")
            
            # Image directory validation
            if 'image_dir' not in local_config:
                errors.append("local mode requires image_dir")
            else:
                image_dir = local_config['image_dir']
                if not os.path.isdir(image_dir):
                    errors.append(f"image_dir does not exist or is not a directory: {image_dir}")
            
            # Output directory validation (optional, will be created if missing)
            if 'output_dir' in local_config:
                output_dir = local_config['output_dir']
                # Check if parent directory exists
                parent_dir = os.path.dirname(os.path.abspath(output_dir))
                if parent_dir and not os.path.isdir(parent_dir):
                    errors.append(f"output_dir parent directory does not exist: {parent_dir}")
    
    elif mode == 'googlecloud':
        if 'googlecloud' not in config:
            errors.append("googlecloud mode requires 'googlecloud' configuration section")
        else:
            gc_config = config['googlecloud']
            
            # Required fields
            required_fields = {
                'project_id': 'project_id',
                'drive_folder_id': 'drive_folder_id',
                'region': 'region',
                'ocr_model_id': 'ocr_model_id',
                'adc_file': 'adc_file'
            }
            
            for field, field_name in required_fields.items():
                if field not in gc_config:
                    errors.append(f"googlecloud mode requires {field_name}")
            
            # ADC file validation
            if 'adc_file' in gc_config:
                adc_file = gc_config['adc_file']
                if not os.path.exists(adc_file):
                    errors.append(f"adc_file does not exist: {adc_file}")
            
            # Drive folder ID format validation (basic check)
            if 'drive_folder_id' in gc_config:
                folder_id = gc_config['drive_folder_id']
                if not isinstance(folder_id, str) or len(folder_id) < 10:
                    errors.append("drive_folder_id appears to be invalid")
            
            # Optional fields with defaults
            if 'max_images' not in config:
                config['max_images'] = 1000
            elif not isinstance(config['max_images'], int) or config['max_images'] < 1:
                errors.append("max_images must be a positive integer")
            
            if 'batch_size_for_doc' not in config:
                config['batch_size_for_doc'] = 10
            elif not isinstance(config['batch_size_for_doc'], int) or config['batch_size_for_doc'] < 1:
                errors.append("batch_size_for_doc must be a positive integer")
    
    return len(errors) == 0, errors


def load_prompt_text(config_or_file: Any) -> str:
    """
    Load prompt text from config or file.
    
    Supports both modes:
    - Legacy: prompt_file (direct file reference)
    - Wizard: prompt_template + context (template assembly)
    
    Args:
        config_or_file: Either a config dict (wizard mode) or prompt file name (legacy mode)
        
    Returns:
        Prompt text as string, or fallback minimal prompt if file not found
    """
    # Check if it's a config dict (wizard mode) or string (legacy mode)
    if isinstance(config_or_file, dict):
        config = config_or_file
        
        # Wizard mode: assemble from template
        if 'prompt_template' in config:
            from wizard.prompt_assembler import PromptAssemblyEngine
            
            assembler = PromptAssemblyEngine()
            context = config.get('context', {})
            template_name = config['prompt_template']
            
            try:
                return assembler.assemble(template_name, context)
            except Exception as e:
                logging.error(f"Failed to assemble prompt from template '{template_name}': {e}")
                raise ValueError(f"Failed to assemble prompt: {e}")
        
        # Legacy mode: use prompt_file
        elif 'prompt_file' in config:
            prompt_file = config['prompt_file']
            return load_prompt_file(prompt_file)
        else:
            raise ValueError("Either prompt_file or prompt_template must be specified in config")
    
    else:
        # Legacy mode: direct file name
        return load_prompt_file(config_or_file)


def load_prompt_file(prompt_file: str) -> str:
    """
    Load prompt text from file in the prompts directory (legacy mode).
    
    Args:
        prompt_file: Name of the prompt file (without path)
        
    Returns:
        Prompt text as string, or fallback minimal prompt if file not found
    """
    prompts_dir = os.path.join(os.path.dirname(__file__), "prompts")
    prompt_path = os.path.join(prompts_dir, prompt_file)
    try:
        with open(prompt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logging.warning(f"Failed to load prompt '{prompt_file}' from {prompt_path}: {e}. Falling back to minimal prompt.")
        return "Transcribe the content of this image as structured text."


def setup_logging(config: dict) -> tuple:
    """
    Set up logging based on configuration.
    
    Args:
        config: Configuration dictionary (normalized, with mode detected)
        
    Returns:
        Tuple of (log_filename, ai_log_filename, ai_logger)
    """
    # Detect mode for log filename
    mode = detect_mode(config)
    
    # Check if logs should be saved to source directory
    save_logs_to_source = False
    if mode == 'local':
        save_logs_to_source = config.get('local', {}).get('save_logs_to_source', False)
    elif mode == 'googlecloud':
        save_logs_to_source = config.get('googlecloud', {}).get('save_logs_to_source', False)
    
    # Determine logs directory
    if save_logs_to_source:
        if mode == 'local':
            # Use image directory for LOCAL mode
            image_dir = config.get('local', {}).get('image_dir', 'local')
            LOGS_DIR = os.path.abspath(image_dir)
        else:
            # For GOOGLECLOUD mode, we'll still use local logs directory initially
            # Logs will be uploaded to Drive after completion
            LOGS_DIR = "logs"
    else:
        # Use configured output_dir for LOCAL mode, or default "logs"
        if mode == 'local':
            LOGS_DIR = config.get('local', {}).get('output_dir', 'logs')
        else:
            LOGS_DIR = "logs"
    
    # Create logs directory if it doesn't exist
    if not os.path.exists(LOGS_DIR):
        os.makedirs(LOGS_DIR)
    
    # Generate timestamp for log files
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Logging setup - use date-time-transcribe-session-{identifier}.log format
    if mode == 'googlecloud':
        # Use drive folder ID for Google Cloud mode
        drive_folder_id = config.get('googlecloud', {}).get('drive_folder_id', 'unknown')
        folder_id_short = drive_folder_id[:8] if len(drive_folder_id) >= 8 else drive_folder_id
        log_filename = os.path.join(LOGS_DIR, f"{timestamp}-transcribe-session-{folder_id_short}.log")
    else:
        # Use image directory name for local mode
        image_dir = config.get('local', {}).get('image_dir', 'local')
        dir_name = os.path.basename(os.path.abspath(image_dir))
        dir_name_short = dir_name[:20] if len(dir_name) > 20 else dir_name
        log_filename = os.path.join(LOGS_DIR, f"{timestamp}-transcribe-session-{dir_name_short}.log")
    
    # Clear any existing handlers to ensure clean logging setup
    # This is important when running from wizard mode where logging might already be configured
    root_logger = logging.getLogger()
    for handler in root_logger.handlers[:]:
        root_logger.removeHandler(handler)
        handler.close()
    
    # Configure logging with file and console handlers
    file_handler = logging.FileHandler(log_filename, encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s'))
    
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s'))
    
    root_logger.setLevel(logging.INFO)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)
    
    # Set up separate logger for AI responses
    ai_logger = logging.getLogger('ai_responses')
    ai_logger.setLevel(logging.INFO)
    ai_log_filename = os.path.join(LOGS_DIR, f"{timestamp}-ai-responses.log")
    ai_handler = logging.FileHandler(ai_log_filename)
    ai_handler.setFormatter(logging.Formatter('%(asctime)s - %(message)s'))
    ai_logger.addHandler(ai_handler)
    ai_logger.propagate = False  # Prevent duplicate logging
    
    return log_filename, ai_log_filename, ai_logger


# ------------------------- MODE ABSTRACTION LAYER - STRATEGY PATTERN -------------------------

class AuthenticationStrategy(ABC):
    """Abstract base class for authentication strategies."""
    
    @abstractmethod
    def authenticate(self) -> Any:
        """
        Authenticate and return credentials/client.
        
        Returns:
            Authentication object (API key string for local, Credentials for googlecloud)
        """
        pass
    
    @abstractmethod
    def validate(self) -> bool:
        """
        Validate authentication is working.
        
        Returns:
            True if authentication is valid, False otherwise
        """
        pass


class LocalAuthStrategy(AuthenticationStrategy):
    """API key authentication for local mode (skeleton implementation)."""
    
    def __init__(self, api_key: str = None):
        """
        Initialize local authentication strategy.
        
        Args:
            api_key: Gemini API key (can also come from GEMINI_API_KEY env var)
        """
        self.api_key = api_key or os.getenv('GEMINI_API_KEY')
        if not self.api_key:
            raise ValueError("API key required (config or GEMINI_API_KEY env var)")
    
    def authenticate(self) -> str:
        """Return API key for Gemini Developer API."""
        return self.api_key
    
    def validate(self) -> bool:
        """Validate API key format (basic check)."""
        return bool(self.api_key and len(self.api_key) > 10)


class GoogleCloudAuthStrategy(AuthenticationStrategy):
    """OAuth2/ADC authentication for Google Cloud mode (skeleton implementation)."""
    
    def __init__(self, adc_file: str):
        """
        Initialize Google Cloud authentication strategy.
        
        Args:
            adc_file: Path to Application Default Credentials file
        """
        self.adc_file = adc_file
    
    def authenticate(self) -> Credentials:
        """Authenticate using ADC file (delegates to existing authenticate() function)."""
        # Will delegate to existing authenticate() function in Phase 4
        return authenticate(self.adc_file)
    
    def validate(self) -> bool:
        """Validate credentials file exists."""
        return os.path.exists(self.adc_file)


class ImageSourceStrategy(ABC):
    """Abstract base class for image source strategies."""
    
    @abstractmethod
    def list_images(self, config: dict) -> list[dict]:
        """
        List available images.
        
        Args:
            config: Configuration dictionary
            
        Returns:
            List of image metadata dictionaries with 'name', 'id', 'webViewLink' keys
        """
        pass
    
    @abstractmethod
    def get_image_bytes(self, image_info: dict) -> bytes:
        """
        Get image bytes.
        
        Args:
            image_info: Image metadata dictionary
            
        Returns:
            Image bytes
        """
        pass
    
    @abstractmethod
    def get_image_url(self, image_info: dict) -> str:
        """
        Get image URL/link for output.
        
        Args:
            image_info: Image metadata dictionary
            
        Returns:
            URL or file path string
        """
        pass


class LocalImageSource(ImageSourceStrategy):
    """Local file system image source."""
    
    def __init__(self, image_dir: str):
        """
        Initialize local image source.
        
        Args:
            image_dir: Path to directory containing images
        """
        self.image_dir = image_dir
        if not os.path.isdir(image_dir):
            raise ValueError(f"Image directory does not exist: {image_dir}")
    
    def list_images(self, config: dict) -> list[dict]:
        """
        List images from local directory with filtering.
        
        Reuses the same filtering logic as Drive-based list_images(),
        supporting all filename patterns and image_start_number/image_count filtering.
        
        Args:
            config: Configuration dictionary
            
        Returns:
            List of image metadata dictionaries with 'name', 'path', 'id', 'webViewLink'
        """
        import glob
        import re
        
        retry_mode = config.get('retry_mode', False)
        retry_image_list = config.get('retry_image_list', [])
        image_start_number = config.get('image_start_number', 1)
        image_count = config.get('image_count', 1000)
        
        # Supported extensions (case-insensitive)
        extensions = ['*.jpg', '*.jpeg', '*.JPG', '*.JPEG']
        all_image_paths = []
        
        for ext in extensions:
            pattern = os.path.join(self.image_dir, ext)
            all_image_paths.extend(glob.glob(pattern))
        
        # Sort by filename
        all_image_paths.sort()
        
        logging.info(f"Found {len(all_image_paths)} total images in local directory (sorted by filename)")
        
        # Convert to dict format compatible with existing code
        all_images = []
        for img_path in all_image_paths:
            filename = os.path.basename(img_path)
            # Normalize path for file:// URL (Windows requires forward slashes)
            abs_path = os.path.abspath(img_path)
            # Convert backslashes to forward slashes for file:// URLs (Windows compatibility)
            normalized_path = abs_path.replace('\\', '/')
            all_images.append({
                'name': filename,
                'path': img_path,
                'id': img_path,  # Use path as ID for local mode
                'webViewLink': f"file://{normalized_path}"  # Local file URL (Windows-compatible)
            })
        
        # RETRY MODE: If enabled, filter for specific failed images only
        if retry_mode:
            logging.info(f"RETRY MODE ENABLED: Looking for {len(retry_image_list)} specific failed images")
            retry_images = []
            
            # Find matching images (exact filename match)
            retry_names = set(retry_image_list)
            for img in all_images:
                if img['name'] in retry_names:
                    retry_images.append(img)
            
            logging.info(f"Found {len(retry_images)} retry images out of {len(retry_image_list)} requested")
            if retry_images:
                retry_filenames = [img['name'] for img in retry_images]
                logging.info(f"Retry images found: {retry_filenames}")
            else:
                logging.warning("No retry images found! Check the retry_image_list names in config.")
            
            return retry_images
        
        # NORMAL MODE: Apply same filtering logic as Drive-based list_images()
        numbered_images = []
        timestamp_images = []
        
        # Helper: case-insensitive check for JPEG extension
        def has_jpeg_extension(filename: str) -> bool:
            lower = filename.lower()
            return lower.endswith('.jpg') or lower.endswith('.jpeg')
        
        # Regex patterns (same as Drive implementation)
        timestamp_pattern = re.compile(r'^image - (\d{4}-\d{2}-\d{2}T\d{6}\.\d{3})\.(?:jpg|jpeg)$', re.IGNORECASE)
        img_date_pattern = re.compile(r'^IMG_\d{8}_(\d+)\.(?:jpg|jpeg)$', re.IGNORECASE)
        
        for img in all_images:
            filename = img['name']
            number = None
            
            # Check for timestamp pattern first
            timestamp_match = timestamp_pattern.match(filename)
            if timestamp_match:
                timestamp_images.append(img)
                continue
            
            # Check for IMG_YYYYMMDD_XXXX.jpg pattern
            img_date_match = img_date_pattern.match(filename)
            if img_date_match:
                try:
                    number = int(img_date_match.group(1))
                except ValueError:
                    continue
            
            # Check if filename matches the pattern image (N).jpg/jpeg
            elif filename.startswith('image (') and (filename.lower().endswith(').jpg') or filename.lower().endswith(').jpeg')):
                try:
                    start_idx = filename.find('(') + 1
                    end_idx = filename.find(')')
                    number_str = filename[start_idx:end_idx]
                    number = int(number_str)
                except (ValueError, IndexError):
                    continue
            
            # Check if filename matches the pattern imageXXXXX.jpg/jpeg
            elif filename.startswith('image') and has_jpeg_extension(filename) and '(' not in filename and ' - ' not in filename and '_' not in filename:
                try:
                    ext_len = 5 if filename.lower().endswith('.jpeg') else 4
                    number_str = filename[5:-ext_len]
                    number = int(number_str)
                except ValueError:
                    continue
            
            # Check if filename matches the pattern XXXXX.jpg/jpeg
            elif has_jpeg_extension(filename) and not filename.startswith('image') and '_' not in filename:
                try:
                    ext_len = 5 if filename.lower().endswith('.jpeg') else 4
                    number_str = filename[:-ext_len]
                    number = int(number_str)
                except ValueError:
                    continue
            
            # Check if filename matches the pattern PREFIX_XXXXX.jpg/jpeg
            elif has_jpeg_extension(filename) and '_' in filename:
                try:
                    ext_len = 5 if filename.lower().endswith('.jpeg') else 4
                    base_no_ext = filename[:-ext_len]
                    underscore_idx = base_no_ext.rfind('_')
                    if underscore_idx != -1:
                        suffix = base_no_ext[underscore_idx + 1:]
                        if suffix.isdigit():
                            number = int(suffix)
                except Exception:
                    continue
            
            # If we found a valid number, check if it's in the desired range
            if number is not None:
                if image_start_number <= number < image_start_number + image_count:
                    numbered_images.append(img)
        
        # Handle numbered images (same logic as Drive implementation)
        filtered_images = []
        
        if numbered_images:
            start_filename_pattern1 = f"image ({image_start_number}).jpg"
            end_filename_pattern1 = f"image ({image_start_number + image_count - 1}).jpg"
            start_filename_pattern2 = f"image{image_start_number:05d}.jpg"
            end_filename_pattern2 = f"image{image_start_number + image_count - 1:05d}.jpg"
            start_filename_pattern3 = f"{image_start_number}.jpg"
            end_filename_pattern3 = f"{image_start_number + image_count - 1}.jpg"
            
            logging.info(f"Filtering numbered images from {start_filename_pattern1} to {end_filename_pattern1} OR {start_filename_pattern2} to {end_filename_pattern2} OR {start_filename_pattern3} to {end_filename_pattern3}")
            
            # Sort numbered images numerically
            def extract_number_for_sorting(img):
                filename = img['name']
                lower = filename.lower()
                if filename.startswith('image (') and (lower.endswith(').jpg') or lower.endswith(').jpeg')):
                    start_idx = filename.find('(') + 1
                    end_idx = filename.find(')')
                    number_str = filename[start_idx:end_idx]
                elif filename.startswith('image') and has_jpeg_extension(filename) and '(' not in filename and ' - ' not in filename and '_' not in filename:
                    ext_len = 5 if lower.endswith('.jpeg') else 4
                    number_str = filename[5:-ext_len]
                elif has_jpeg_extension(filename) and '_' in filename:
                    ext_len = 5 if lower.endswith('.jpeg') else 4
                    base_no_ext = filename[:-ext_len]
                    underscore_idx = base_no_ext.rfind('_')
                    number_str = base_no_ext[underscore_idx + 1:]
                else:
                    ext_len = 5 if lower.endswith('.jpeg') else 4
                    number_str = filename[:-ext_len]
                return int(number_str)
            
            numbered_images.sort(key=extract_number_for_sorting)
            filtered_images.extend(numbered_images)
        
        # Handle timestamp images
        if timestamp_images:
            logging.info(f"Found {len(timestamp_images)} timestamp-based images")
            
            # Sort timestamp images chronologically
            def extract_timestamp_for_sorting(img):
                filename = img['name']
                match = timestamp_pattern.match(filename)
                if match:
                    timestamp_str = match.group(1)
                    try:
                        from datetime import datetime
                        formatted_timestamp = f"{timestamp_str[:11]}{timestamp_str[11:13]}:{timestamp_str[13:15]}:{timestamp_str[15:]}"
                        return datetime.fromisoformat(formatted_timestamp)
                    except ValueError:
                        return datetime.min
                return datetime.min
            
            timestamp_images.sort(key=extract_timestamp_for_sorting)
            
            # For timestamp images, treat image_start_number as starting position
            start_pos = max(1, image_start_number) - 1
            end_pos = min(len(timestamp_images), start_pos + image_count)
            
            selected_timestamp_images = timestamp_images[start_pos:end_pos]
            filtered_images.extend(selected_timestamp_images)
            
            if selected_timestamp_images:
                logging.info(f"Selected {len(selected_timestamp_images)} timestamp images from position {image_start_number} to {start_pos + len(selected_timestamp_images)}")
        
        # Fallback: if no images selected, use position-based selection
        if not filtered_images and all_images:
            start_pos = max(1, image_start_number) - 1
            end_pos = min(len(all_images), start_pos + image_count)
            fallback_selected = all_images[start_pos:end_pos]
            if fallback_selected:
                logging.info(f"No numeric/timestamp matches; falling back to position selection: items {image_start_number} to {image_start_number + len(fallback_selected) - 1}")
                filtered_images = fallback_selected
        
        logging.info(f"Selected {len(filtered_images)} total images for processing")
        
        if filtered_images:
            filenames = [img['name'] for img in filtered_images]
            logging.info(f"Final selected files: {filenames}")
        
        return filtered_images
    
    def get_image_bytes(self, image_info: dict) -> bytes:
        """
        Read image from local file system.
        
        Args:
            image_info: Image metadata dictionary with 'path' key
            
        Returns:
            Image bytes
        """
        with open(image_info['path'], 'rb') as f:
            return f.read()
    
    def get_image_url(self, image_info: dict) -> str:
        """
        Return local file path with normalized separators.
        
        Args:
            image_info: Image metadata dictionary with 'path' key
            
        Returns:
            Local file path with forward slashes (cross-platform compatible)
        """
        # Normalize path separators to forward slashes for consistency across platforms
        path = image_info['path']
        return path.replace('\\', '/')


class DriveImageSource(ImageSourceStrategy):
    """Google Drive image source."""
    
    def __init__(self, drive_service, drive_folder_id: str, document_name: str = "Unknown"):
        """
        Initialize Drive image source.
        
        Args:
            drive_service: Google Drive API service
            drive_folder_id: Drive folder ID containing images
            document_name: Document name for logging (optional)
        """
        self.drive_service = drive_service
        self.drive_folder_id = drive_folder_id
        self.document_name = document_name
    
    def list_images(self, config: dict) -> list[dict]:
        """
        List images from Google Drive (delegates to existing function).
        
        Args:
            config: Configuration dictionary
            
        Returns:
            List of image metadata dictionaries
        """
        return list_images(self.drive_service, config)
    
    def get_image_bytes(self, image_info: dict) -> bytes:
        """
        Download image from Drive (delegates to existing function).
        
        Args:
            image_info: Image metadata dictionary with 'id' and 'name'
            
        Returns:
            Image bytes
        """
        return download_image(
            self.drive_service,
            image_info['id'],
            image_info['name'],
            self.document_name
        )
    
    def get_image_url(self, image_info: dict) -> str:
        """
        Return Drive web view link.
        
        Args:
            image_info: Image metadata dictionary
            
        Returns:
            Web view link
        """
        return image_info.get('webViewLink', '')


class AIClientStrategy(ABC):
    """Abstract base class for AI client strategies."""
    
    @abstractmethod
    def transcribe(self, image_bytes: bytes, filename: str, prompt: str) -> tuple[str, float, dict]:
        """
        Transcribe image using AI model.
        
        Args:
            image_bytes: Image file bytes
            filename: Image filename (for logging)
            prompt: Transcription prompt text
            
        Returns:
            Tuple of (transcription_text, elapsed_time, usage_metadata)
        """
        pass


class GeminiDevClient(AIClientStrategy):
    """Gemini Developer API client."""
    
    def __init__(self, api_key: str, model_id: str = "gemini-3-flash-preview", ai_logger=None):
        """
        Initialize Gemini Developer API client.
        
        Args:
            api_key: Gemini API key
            model_id: Model ID to use (default: gemini-3-flash-preview)
            ai_logger: Logger instance for AI responses (optional)
        """
        self.api_key = api_key
        self.model_id = model_id
        self.ai_logger = ai_logger
        # Initialize Gemini client for Developer API (not Vertex AI)
        self.client = genai.Client(api_key=api_key)
        logging.info(f"Gemini Developer API client initialized with model {model_id}")
    
    def transcribe(self, image_bytes: bytes, filename: str, prompt: str) -> tuple[str, float, dict]:
        """
        Transcribe image using Gemini Developer API.
        
        Args:
            image_bytes: Image file bytes
            filename: Image filename (for logging)
            prompt: Transcription prompt text
            
        Returns:
            Tuple of (transcription_text, elapsed_time, usage_metadata)
        """
        import time
        
        function_start_time = time.time()
        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Starting transcription for image '{filename}' (size: {len(image_bytes)} bytes)")
        
        # Create image part
        image_part = types.Part.from_bytes(
            data=image_bytes,
            mime_type="image/jpeg"
        )
        
        # Create content with prompt and image
        content = types.Content(
            role="user",
            parts=[
                types.Part.from_text(text=prompt),
                image_part
            ]
        )
        
        # Configure generation parameters (matching Vertex AI configuration)
        generate_content_config = types.GenerateContentConfig(
            temperature=0.1,
            top_p=0.8,
            seed=0,
            max_output_tokens=65535,
            system_instruction=[types.Part.from_text(text=prompt)],
            thinking_config=types.ThinkingConfig(
                thinking_budget=5000,
            ),
        )
        
        max_retries = 3
        retry_delay = 30  # seconds
        timeout_seconds_list = [60, 120, 300]  # 1 min, 2 min, 5 min
        
        for attempt in range(max_retries):
            attempt_start_time = time.time()
            timeout_seconds = timeout_seconds_list[attempt]
            
            try:
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} for image '{filename}' (timeout: {timeout_seconds/60:.1f} min)")
                
                # Make API call
                api_call_start = time.time()
                response = self.client.models.generate_content(
                    model=self.model_id,
                    contents=[content],
                    config=generate_content_config
                )
                
                api_call_elapsed = time.time() - api_call_start
                elapsed_time = time.time() - attempt_start_time
                total_elapsed = time.time() - function_start_time
                
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Gemini API response received in {api_call_elapsed:.1f}s (attempt total: {elapsed_time:.1f}s, function total: {total_elapsed:.1f}s) for '{filename}'")
                
                # Log warning if API call took unusually long
                if api_call_elapsed > 60:
                    logging.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: API call took {api_call_elapsed:.1f}s (>60s) for '{filename}'")
                
                # Extract text from response - try response.text first, then candidates
                text = None
                if hasattr(response, 'text') and response.text:
                    text = response.text
                elif hasattr(response, 'candidates') and response.candidates:
                    # Try to extract text from first candidate
                    candidate = response.candidates[0]
                    if hasattr(candidate, 'content') and candidate.content:
                        if hasattr(candidate.content, 'parts') and candidate.content.parts:
                            # Extract text from parts
                            text_parts = []
                            for part in candidate.content.parts:
                                if hasattr(part, 'text') and part.text:
                                    text_parts.append(part.text)
                            if text_parts:
                                text = '\n'.join(text_parts)
                
                if not text:
                    text = "[No transcription text received]"
                
                # Extract usage metadata if available
                usage_metadata = {}
                if hasattr(response, 'usage_metadata'):
                    usage_metadata = {
                        'prompt_tokens': getattr(response.usage_metadata, 'prompt_token_count', 0),
                        'completion_tokens': getattr(response.usage_metadata, 'candidates_token_count', 0),
                        'total_tokens': getattr(response.usage_metadata, 'total_token_count', 0),
                        'cached_tokens': getattr(response.usage_metadata, 'cached_content_token_count', 0)
                    }
                
                # Log the full AI response to the AI responses log (similar to transcribe_image)
                if self.ai_logger:
                    self.ai_logger.info(f"=== AI Response for {filename} ===")
                    self.ai_logger.info(f"Model: {self.model_id}")
                    self.ai_logger.info(f"Request timestamp: {datetime.now().isoformat()}")
                    self.ai_logger.info(f"Image size: {len(image_bytes)} bytes")
                    self.ai_logger.info(f"Prompt length: {len(prompt)} characters")
                    self.ai_logger.info(f"Response length: {len(text) if text else 0} characters")
                    self.ai_logger.info(f"Processing time: {elapsed_time:.1f} seconds")
                    
                    # Log response metadata if available
                    if hasattr(response, 'usage_metadata'):
                        self.ai_logger.info(f"Usage metadata: {response.usage_metadata}")
                    if hasattr(response, 'candidates') and response.candidates:
                        self.ai_logger.info(f"Number of candidates: {len(response.candidates)}")
                        if hasattr(response.candidates[0], 'finish_reason'):
                            self.ai_logger.info(f"Finish reason: {response.candidates[0].finish_reason}")
                    
                    # Log full response text (prompt is logged only once at session start)
                    self.ai_logger.info(f"Full response:\n{text}")
                    self.ai_logger.info(f"=== End AI Response for {filename} ===\n")
                
                function_total_elapsed = time.time() - function_start_time
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Transcription completed for '{filename}' in {function_total_elapsed:.1f}s total")
                
                return text, elapsed_time, usage_metadata
                
            except ServerError as e:
                # Handle server errors (503, 500, etc.) - retry with exponential backoff
                attempt_elapsed = time.time() - attempt_start_time
                total_elapsed = time.time() - function_start_time
                error_type = type(e).__name__
                
                # Get status code from exception
                status_code = getattr(e, 'status_code', None)
                error_str = str(e)
                
                # If status_code not available, extract from exception string (format: "503 UNAVAILABLE")
                if status_code is None:
                    import re
                    match = re.match(r'(\d+)', error_str)
                    if match:
                        status_code = int(match.group(1))
                    else:
                        status_code = None
                
                # Retry 503 (Service Unavailable) and 500 (Internal Server Error) errors
                is_retryable = status_code in (503, 500) if status_code else True  # Default to retryable if unknown
                
                error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} failed for '{filename}' after {attempt_elapsed:.1f}s (total: {total_elapsed:.1f}s): {error_type} (status {status_code}): {str(e)}"
                logging.warning(error_msg)
                
                if is_retryable and attempt < max_retries - 1:
                    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Retrying in {retry_delay}s... (exponential backoff)")
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                else:
                    if not is_retryable:
                        logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] Non-retryable server error (status {status_code}) for '{filename}' after {attempt_elapsed:.1f}s: {error_str}")
                        logging.error(f"Full traceback:\n{traceback.format_exc()}")
                        raise RuntimeError(f"Server error (status {status_code}): {error_str}") from e
                    else:
                        logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] All {max_retries} attempts failed for '{filename}' after {total_elapsed:.1f}s: {error_type} (status {status_code}): {str(e)}")
                        return f"[Error during transcription: {str(e)}]", None, None
                    
            except (TimeoutError, ConnectionError, OSError) as e:
                attempt_elapsed = time.time() - attempt_start_time
                total_elapsed = time.time() - function_start_time
                error_type = type(e).__name__
                
                error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} failed for '{filename}' after {attempt_elapsed:.1f}s (total: {total_elapsed:.1f}s): {error_type}: {str(e)}"
                logging.warning(error_msg)
                
                if attempt < max_retries - 1:
                    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Retrying in {retry_delay}s... (exponential backoff)")
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                else:
                    logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] All {max_retries} attempts failed for '{filename}' after {total_elapsed:.1f}s: {error_type}: {str(e)}")
                    return f"[Error during transcription: {str(e)}]", None, None
                    
            except ClientError as e:
                # Handle API errors (400, 401, 403, etc.) - don't retry, fail immediately
                attempt_elapsed = time.time() - attempt_start_time
                total_elapsed = time.time() - function_start_time
                
                # Get status code from exception - ClientError has status_code attribute
                status_code = getattr(e, 'status_code', None)
                error_str = str(e)
                
                # If status_code not available, extract from exception string (format: "400 INVALID_ARGUMENT")
                if status_code is None:
                    # Parse from string like "400 INVALID_ARGUMENT. {...}"
                    import re
                    match = re.match(r'(\d+)', error_str)
                    if match:
                        status_code = int(match.group(1))
                    else:
                        status_code = None
                
                # Check for API key errors: status 400 with API key indicators
                # Primary check: status code 400
                # Secondary check: error message contains API key related text
                is_api_key_error = False
                if status_code == 400:
                    # Check error message/reason for API key issues
                    if 'API key' in error_str or 'API_KEY' in error_str or 'API_KEY_INVALID' in error_str:
                        is_api_key_error = True
                    # Also check response JSON if available
                    elif hasattr(e, 'response') and isinstance(e.response, dict):
                        error_info = e.response.get('error', {})
                        if isinstance(error_info, dict):
                            details = error_info.get('details', [])
                            for detail in details:
                                if isinstance(detail, dict) and detail.get('reason') == 'API_KEY_INVALID':
                                    is_api_key_error = True
                                    break
                
                if is_api_key_error:
                    error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Invalid API key error (status 400) for '{filename}' after {attempt_elapsed:.1f}s"
                    logging.error(error_msg)
                    logging.error(f"Error details: {error_str}")
                    logging.error(f"Full traceback:\n{traceback.format_exc()}")
                    # Raise immediately - don't return error string, let it propagate
                    raise ValueError(f"Invalid API key (status 400): {error_str}") from e
                
                # For other ClientErrors with status codes (400, 401, 403, etc.), don't retry - fail immediately
                if status_code is not None:
                    error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] API error (status {status_code}) for '{filename}' after {attempt_elapsed:.1f}s: {error_str}"
                    logging.error(error_msg)
                    logging.error(f"Full traceback:\n{traceback.format_exc()}")
                    # Raise immediately for non-retryable errors
                    raise RuntimeError(f"API error (status {status_code}): {error_str}") from e
                else:
                    # If we can't determine status code, treat as unexpected error
                    error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] API error for '{filename}' after {attempt_elapsed:.1f}s: {error_str}"
                    logging.error(error_msg)
                    logging.error(f"Full traceback:\n{traceback.format_exc()}")
                    raise RuntimeError(f"API error: {error_str}") from e
                
            except Exception as e:
                attempt_elapsed = time.time() - attempt_start_time
                total_elapsed = time.time() - function_start_time
                error_type = type(e).__name__
                
                error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Unexpected error in Gemini API transcription for '{filename}' after {attempt_elapsed:.1f}s: {error_type}: {str(e)}"
                logging.error(error_msg)
                logging.error(f"Full traceback:\n{traceback.format_exc()}")
                return f"[Error during transcription: {str(e)}]", None, None


class VertexAIClient(AIClientStrategy):
    """Vertex AI client (skeleton implementation)."""
    
    def __init__(self, genai_client, model_id: str):
        """
        Initialize Vertex AI client.
        
        Args:
            genai_client: Initialized genai.Client
            model_id: Model ID to use
        """
        self.genai_client = genai_client
        self.model_id = model_id
    
    def transcribe(self, image_bytes: bytes, filename: str, prompt: str) -> tuple[str, float, dict]:
        """Transcribe using Vertex AI (delegates to existing function)."""
        # Will delegate to existing transcribe_image() function in Phase 4
        return transcribe_image(self.genai_client, image_bytes, filename, prompt, self.model_id)


class OutputStrategy(ABC):
    """Abstract base class for output strategies."""
    
    @abstractmethod
    def initialize(self, config: dict) -> Any:
        """
        Initialize output destination.
        
        Args:
            config: Configuration dictionary
            
        Returns:
            Output identifier (doc_id for Google Docs, log_path for local)
        """
        pass
    
    @abstractmethod
    def write_batch(self, pages: list[dict], batch_num: int, is_first: bool) -> None:
        """
        Write batch of transcriptions.
        
        Args:
            pages: List of page dictionaries with transcription data
            batch_num: Batch number (1-based)
            is_first: True if this is the first batch
        """
        pass
    
    @abstractmethod
    def finalize(self, all_pages: list[dict], metrics: dict, start_time=None, end_time=None, error_info=None) -> None:
        """
        Finalize output (update overview, close files, etc.).
        
        Args:
            all_pages: All transcribed pages
            metrics: Session metrics dictionary
            start_time: Session start datetime
            end_time: Session end datetime
            error_info: Optional dict with error details {'type': str, 'message': str, 'status_code': int, 'next_image_number': int}
        """
        pass


class LogFileOutput(OutputStrategy):
    """Log file output for local mode."""
    
    def __init__(self, output_dir: str, ai_logger):
        """
        Initialize log file output.
        
        Args:
            output_dir: Directory for output log files
            ai_logger: Logger instance for AI responses
        """
        self.output_dir = output_dir
        self.ai_logger = ai_logger
        self.log_file_path = None
        self.written_images = set()  # Track written images to prevent duplicates
        os.makedirs(output_dir, exist_ok=True)
    
    def initialize(self, config: dict, prompt_text: str = None) -> str:
        """
        Initialize log file for transcription output.
        
        Creates a timestamped log file and writes session metadata.
        
        Args:
            config: Configuration dictionary
            prompt_text: Full prompt text used for transcription (optional)
            
        Returns:
            Log file path
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        archive_index = config.get('archive_index', 'transcription')
        self.log_file_path = os.path.join(
            self.output_dir,
            f"{timestamp}-{archive_index}-transcription.log"
        )
        
        # Write session header to log file
        with open(self.log_file_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("TRANSCRIPTION SESSION - LOCAL MODE\n")
            f.write("=" * 80 + "\n")
            f.write(f"Session started: {datetime.now().isoformat()}\n")
            f.write(f"Archive index: {config.get('archive_index', 'N/A')}\n")
            f.write(f"Image directory: {config.get('local', {}).get('image_dir', 'N/A')}\n")
            f.write(f"Model: {config.get('local', {}).get('ocr_model_id', 'N/A')}\n")
            f.write(f"Image range: {config.get('image_start_number', 'N/A')} to {config.get('image_start_number', 0) + config.get('image_count', 0) - 1}\n")
            f.write("=" * 80 + "\n")
            
            # Add full prompt text if provided
            if prompt_text:
                f.write("\n" + "=" * 80 + "\n")
                f.write("PROMPT TEXT USED FOR TRANSCRIPTION\n")
                f.write("=" * 80 + "\n")
                f.write(f"{prompt_text}\n")
                f.write("=" * 80 + "\n\n")
            
            f.write("\n")
        
        logging.info(f"Created transcription log file: {self.log_file_path}")
        return self.log_file_path
    
    def write_batch(self, pages: list[dict], batch_num: int, is_first: bool) -> None:
        """
        Write batch of transcriptions to log file.
        
        Args:
            pages: List of page dictionaries with transcription data
            batch_num: Batch number (1-based)
            is_first: True if this is the first batch
        """
        if not self.log_file_path:
            raise ValueError("Log file not initialized. Call initialize() first.")
        
        # Filter out already-written images to prevent duplicates
        pages_to_write = []
        for page in pages:
            page_id = f"{page['name']}_{batch_num}"
            if page_id not in self.written_images:
                pages_to_write.append(page)
                self.written_images.add(page_id)
            else:
                logging.warning(f"Skipping duplicate write for {page['name']} (batch {batch_num})")
        
        if not pages_to_write:
            return  # Nothing new to write
        
        with open(self.log_file_path, 'a', encoding='utf-8') as f:
            if is_first:
                f.write("\n" + "=" * 80 + "\n")
                f.write("TRANSCRIPTIONS\n")
                f.write("=" * 80 + "\n\n")
            
            for page in pages_to_write:
                f.write("-" * 80 + "\n")
                f.write(f"Image: {page['name']}\n")
                f.write(f"Source: {page.get('webViewLink', page.get('path', ''))}\n")
                f.write("-" * 80 + "\n")
                f.write(f"{page['text']}\n")
                f.write("\n")
        
        logging.info(f"Wrote batch {batch_num} ({len(pages_to_write)} pages) to log file")
    
    def finalize(self, all_pages: list[dict], metrics: dict, start_time=None, end_time=None, error_info=None) -> None:
        """
        Finalize log file with session summary.
        
        Args:
            all_pages: All transcribed pages
            metrics: Session metrics dictionary
            start_time: Session start datetime
            end_time: Session end datetime
            error_info: Optional dict with error details
        """
        if not self.log_file_path:
            raise ValueError("Log file not initialized. Call initialize() first.")
        
        with open(self.log_file_path, 'a', encoding='utf-8') as f:
            f.write("\n" + "=" * 80 + "\n")
            f.write("SESSION SUMMARY\n")
            f.write("=" * 80 + "\n")
            
            # Add error information if present
            if error_info:
                f.write("**RUN INTERRUPTED BY ERROR**\n")
                f.write(f"Error Type: {error_info.get('type', 'Unknown')}\n")
                if error_info.get('status_code'):
                    f.write(f"Error Code: {error_info['status_code']}\n")
                f.write(f"Error Message: {error_info.get('message', 'Unknown error')}\n")
                if error_info.get('next_image_number'):
                    f.write(f"\nTo resume: Update config 'image_start_number' to {error_info['next_image_number']}\n")
                f.write("\n")
            
            f.write(f"Session completed: {datetime.now().isoformat()}\n")
            f.write(f"Total images processed: {len(all_pages)}\n")
            f.write(f"Successful transcriptions: {len([p for p in all_pages if p.get('text') and not p['text'].startswith('[Error')])}\n")
            f.write(f"Failed transcriptions: {len([p for p in all_pages if not p.get('text') or p['text'].startswith('[Error')])}\n")
            
            if metrics:
                f.write("\nMetrics:\n")
                for key, value in metrics.items():
                    f.write(f"  {key}: {value}\n")
            
            f.write("=" * 80 + "\n")
        
        logging.info(f"Finalized transcription log file: {self.log_file_path}")


class GoogleDocsOutput(OutputStrategy):
    """Google Docs output."""
    
    def __init__(self, docs_service, drive_service, genai_client, config: dict, prompt_text: str):
        """
        Initialize Google Docs output.
        
        Args:
            docs_service: Google Docs API service
            drive_service: Google Drive API service
            genai_client: Vertex AI Gemini client
            config: Configuration dictionary
            prompt_text: Transcription prompt text
        """
        self.docs_service = docs_service
        self.drive_service = drive_service
        self.genai_client = genai_client
        self.config = config
        self.prompt_text = prompt_text
        self.doc_id = None
        self.start_time = None
        self.end_time = None
    
    def initialize(self, config: dict) -> str:
        """
        Create Google Doc (delegates to existing function).
        
        Args:
            config: Configuration dictionary
            
        Returns:
            Document ID
        """
        run_date = datetime.now().strftime("%Y%m%d")
        # Handle both normalized (nested) and legacy (flat) config formats
        googlecloud_config = config.get('googlecloud', {})
        document_name = googlecloud_config.get('document_name') or config.get('document_name', 'Unknown')
        doc_name = f"{document_name} {run_date}"
        self.doc_id = create_doc(self.docs_service, self.drive_service, doc_name, config)
        self.start_time = datetime.now()
        logging.info(f"Created Google Doc with ID: {self.doc_id}")
        return self.doc_id
    
    def write_batch(self, pages: list[dict], batch_num: int, is_first: bool) -> None:
        """
        Write batch to Google Doc (delegates to existing write_to_doc()).
        
        Args:
            pages: List of ALL page dictionaries with transcription data (accumulated so far)
            batch_num: Batch number (1-based)
            is_first: True if this is the first batch
        """
        if not self.doc_id:
            raise ValueError("Document not initialized. Call initialize() first.")
        
        # Calculate start index (number of pages already written)
        batch_size = self.config.get('batch_size_for_doc', 10)
        if is_first:
            start_idx = 0
        else:
            # For subsequent batches, start_idx is the number of pages in previous batches
            start_idx = (batch_num - 1) * batch_size
        
        # For first batch, include overview; for subsequent batches, skip overview
        write_overview = is_first
        
        # Delegate to existing write_to_doc() function
        # Note: write_to_doc expects all pages and will write from start_idx onwards
        write_to_doc(
            self.docs_service,
            self.drive_service,
            self.doc_id,
            pages,
            self.config,
            self.prompt_text,
            start_idx=start_idx,
            metrics=None,  # Metrics will be updated in finalize()
            start_time=self.start_time,
            end_time=None,  # End time not known yet
            write_overview=write_overview,
            genai_client=self.genai_client
        )
        
        logging.info(f"Wrote batch {batch_num} (pages {start_idx} onwards) to Google Doc")
    
    def finalize(self, all_pages: list[dict], metrics: dict, start_time=None, end_time=None, error_info=None) -> None:
        """
        Update overview section (delegates to existing update_overview_section()).
        
        Args:
            all_pages: All transcribed pages
            metrics: Session metrics dictionary
            start_time: Session start datetime (overrides self.start_time if provided)
            end_time: Session end datetime (overrides self.end_time if provided)
            error_info: Optional dict with error details (not used for Google Docs mode)
        """
        if not self.doc_id:
            raise ValueError("Document not initialized. Call initialize() first.")
        
        # Use provided times or fall back to instance times
        final_start_time = start_time or self.start_time
        final_end_time = end_time or self.end_time or datetime.now()
        
        # Delegate to existing update_overview_section() function
        # Note: update_overview_section doesn't currently support error_info, but that's OK for Google Cloud mode
        update_overview_section(
            self.docs_service,
            self.doc_id,
            all_pages,
            self.config,
            self.prompt_text,
            metrics=metrics,
            start_time=final_start_time,
            end_time=final_end_time
        )
        
        logging.info(f"Finalized Google Doc with overview update")

class CompositeOutput(OutputStrategy):
    """Composite output strategy that delegates to multiple output strategies."""
    
    def __init__(self, strategies: list):
        """
        Initialize composite output with multiple strategies.
        
        Args:
            strategies: List of OutputStrategy instances
        """
        self.strategies = strategies
    
    def initialize(self, config: dict, prompt_text: str = None) -> list:
        """Initialize all strategies and return list of output IDs."""
        output_ids = []
        for strategy in self.strategies:
            try:
                output_id = strategy.initialize(config, prompt_text)
                output_ids.append(output_id)
            except Exception as e:
                logging.error(f"Error initializing output strategy {strategy.__class__.__name__}: {e}")
        return output_ids
    
    def write_batch(self, pages: list[dict], batch_num: int, is_first: bool) -> None:
        """Write batch to all strategies."""
        if not pages:
            logging.warning(f"CompositeOutput.write_batch called with empty pages list (batch {batch_num})")
        
        for strategy in self.strategies:
            try:
                strategy.write_batch(pages, batch_num, is_first)
            except Exception as e:
                logging.error(f"Error writing batch to {strategy.__class__.__name__}: {e}")
                import traceback
                logging.error(traceback.format_exc())
    
    def finalize(self, all_pages: list[dict], metrics: dict, start_time=None, end_time=None, error_info=None) -> None:
        """Finalize all strategies."""
        for strategy in self.strategies:
            try:
                strategy.finalize(all_pages, metrics, start_time, end_time, error_info)
            except Exception as e:
                logging.error(f"Error finalizing {strategy.__class__.__name__}: {e}")


class MarkdownOutput(OutputStrategy):
    """Markdown output for local mode."""
    
    def __init__(self, target_dir: str, config: dict, prompt_text: str):
        """
        Initialize Markdown output.
        
        Args:
            target_dir: Directory where markdown file will be saved (image_dir)
            config: Configuration dictionary
            prompt_text: Transcription prompt text
        """
        self.target_dir = target_dir
        self.config = config
        self.prompt_text = prompt_text
        self.temp_body_file = None
        self.final_file_path = None
    
    def initialize(self, config: dict, prompt_text: str = None) -> str:
        """Initialize markdown file."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        archive_index = config.get('archive_index', 'transcription')
        self.final_file_path = os.path.join(self.target_dir, f"{archive_index}_{timestamp}.md")
        self.temp_body_file = os.path.join(self.target_dir, f"temp_body_{timestamp}.md")
        
        # Create empty temp file
        open(self.temp_body_file, 'w', encoding='utf-8').close()
        
        logging.info(f"Initialized Markdown output: {self.final_file_path}")
        return self.final_file_path
    
    def write_batch(self, pages: list[dict], batch_num: int, is_first: bool) -> None:
        """Write batch to temp markdown file."""
        if not self.temp_body_file:
            raise ValueError("Markdown output not initialized")
        
        if not pages:
            logging.warning(f"MarkdownOutput.write_batch called with empty pages list (batch {batch_num})")
            return
        
        try:
            with open(self.temp_body_file, 'a', encoding='utf-8') as f:
                for page in pages:
                    if not page:
                        logging.warning(f"MarkdownOutput.write_batch: skipping empty page in batch {batch_num}")
                        continue
                    
                    page_name = page.get('name', 'Unknown')
                    f.write(f"\n---\n\n## {page_name}\n\n")
                    link = page.get('webViewLink', page.get('path', ''))
                    if link:
                        # For local files, use relative path (just filename) since markdown is in same directory
                        # This prevents path duplication when markdown is opened in browser/viewer
                        # Markdown file is saved in target_dir (image_dir), same as images
                        if link.startswith('http://') or link.startswith('https://'):
                            # Keep Google Drive URLs as-is (GOOGLECLOUD mode)
                            pass
                        else:
                            # For local files (file:// URLs, absolute paths, or relative paths),
                            # use just the filename since markdown and images are in the same directory
                            link = page_name
                        f.write(f"**Source:** [{page_name}]({link})\n\n")
                    
                    # Preserve newlines by adding two spaces before newlines for markdown line breaks
                    text = page.get('text', '')
                    if not text:
                        logging.warning(f"MarkdownOutput.write_batch: page '{page_name}' has empty text field")
                        text = "[No transcription text available]"
                    
                    # Replace single newlines with two spaces + newline (markdown line break)
                    # but preserve paragraph breaks (double newlines)
                    text = text.replace('\n\n', '\n\n')  # Preserve paragraph breaks
                    text = text.replace('\n', '  \n')  # Convert single newlines to markdown line breaks
                    f.write(f"{text}\n")
            
            logging.debug(f"MarkdownOutput.write_batch: wrote {len(pages)} page(s) to temp file {self.temp_body_file}")
        except Exception as e:
            logging.error(f"MarkdownOutput.write_batch failed for batch {batch_num}: {e}")
            import traceback
            logging.error(traceback.format_exc())
            raise
    
    def finalize(self, all_pages: list[dict], metrics: dict, start_time=None, end_time=None, error_info=None) -> None:
        """Finalize markdown file with overview at top."""
        if not self.final_file_path:
            raise ValueError("Markdown output not initialized")
        
        # Generate overview
        overview = create_local_overview_section(
            all_pages, self.config, self.prompt_text, 
            metrics=metrics, start_time=start_time, end_time=end_time, error_info=error_info
        )
        
        # Write final file
        with open(self.final_file_path, 'w', encoding='utf-8') as final:
            archive_index = self.config.get('archive_index', 'Transcription')
            final.write(f"# {archive_index}\n\n")
            final.write(f"## Session Overview\n\n```text\n{overview}\n```\n\n")
            
            # Append temp body content
            if os.path.exists(self.temp_body_file):
                with open(self.temp_body_file, 'r', encoding='utf-8') as temp:
                    final.write(temp.read())
                os.remove(self.temp_body_file)
        
        logging.info(f"Finalized Markdown output: {self.final_file_path}")






class WordOutput(OutputStrategy):
    """Microsoft Word output for local mode."""
    
    def __init__(self, target_dir: str, config: dict, prompt_text: str):
        """Initialize Word output."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word output. Install with: pip install python-docx>=0.8.11")
        self.target_dir = target_dir
        self.config = config
        self.prompt_text = prompt_text
        self.doc = None
        self.doc_path = None
        self.overview_placeholder = None
    
    def initialize(self, config: dict, prompt_text: str = None) -> str:
        """Initialize Word document."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        archive_index = config.get('archive_index', 'transcription')
        self.doc_path = os.path.join(self.target_dir, f"{archive_index}_{timestamp}.docx")
        self.doc = Document()
        self.doc.add_heading(archive_index, level=1)
        self.overview_placeholder = self.doc.add_paragraph("[OVERVIEW_PLACEHOLDER]")
        self.doc.save(self.doc_path)
        logging.info(f"Initialized Word output: {self.doc_path}")
        return self.doc_path
    
    def write_batch(self, pages: list[dict], batch_num: int, is_first: bool) -> None:
        """Write batch to Word document."""
        if not self.doc:
            raise ValueError("Word output not initialized")
        
        if not pages:
            logging.warning(f"WordOutput.write_batch called with empty pages list (batch {batch_num})")
            return
        
        archive_index = self.config.get('archive_index', '')
        
        try:
            for idx, page in enumerate(pages, start=1):
                if not page:
                    logging.warning(f"WordOutput.write_batch: skipping empty page in batch {batch_num}")
                    continue
                
                # Calculate page number (cumulative across batches)
                page_number = (batch_num - 1) + idx if batch_num > 0 else idx
                
                page_name = page.get('name', 'Unknown')
                # Add page header like "ф487оп1спр545стр1" (matching Google Docs format)
                if archive_index:
                    page_header = f"{archive_index}стр{page_number}"
                else:
                    page_header = page_name
                
                self.doc.add_heading(page_header, level=2)
                
                # Add source link
                link = page.get('webViewLink', page.get('path', ''))
                if link:
                    p = self.doc.add_paragraph()
                    p.add_run('Src Img Url: ').bold = True
                    p.add_run(page_name)
                    self.doc.add_paragraph()
                
                # Add transcription text with bold formatting for **text**
                text = page.get('text', '')
                if not text:
                    logging.warning(f"WordOutput.write_batch: page '{page_name}' has empty text field")
                    text = "[No transcription text available]"
                
                self._add_formatted_text(text)
                
                # Add separator
                self.doc.add_paragraph('_' * 80)
            
            self.doc.save(self.doc_path)
            logging.debug(f"WordOutput.write_batch: wrote {len(pages)} page(s) to Word document {self.doc_path}")
        except Exception as e:
            logging.error(f"WordOutput.write_batch failed for batch {batch_num}: {e}")
            import traceback
            logging.error(traceback.format_exc())
            raise
    
    def _add_formatted_text(self, text: str) -> None:
        """
        Add text to document with bold formatting for **text** patterns.
        
        Args:
            text: Text with markdown-style bold markers
        """
        import re
        
        # Split text by lines to preserve paragraph structure
        lines = text.split('\n')
        
        for line in lines:
            if not line.strip():
                self.doc.add_paragraph()  # Empty line
                continue
            
            p = self.doc.add_paragraph()
            
            # Pattern to match **text**
            pattern = r'\*\*(.+?)\*\*'
            last_end = 0
            
            for match in re.finditer(pattern, line):
                # Add text before the match (normal)
                if match.start() > last_end:
                    p.add_run(line[last_end:match.start()])
                
                # Add the matched text (bold, without the **)
                bold_run = p.add_run(match.group(1))
                bold_run.bold = True
                
                last_end = match.end()
            
            # Add remaining text after last match
            if last_end < len(line):
                p.add_run(line[last_end:])
    
    def finalize(self, all_pages: list[dict], metrics: dict, start_time=None, end_time=None, error_info=None) -> None:
        """Finalize Word document with overview."""
        if not self.doc or not self.overview_placeholder:
            raise ValueError("Word output not initialized")
        overview = create_local_overview_section(all_pages, self.config, self.prompt_text, metrics=metrics, start_time=start_time, end_time=end_time, error_info=error_info)
        self.overview_placeholder.clear()
        run = self.overview_placeholder.add_run(overview)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        self.doc.save(self.doc_path)
        logging.info(f"Finalized Word output: {self.doc_path}")


class ModeFactory:
    """Factory for creating mode-specific components."""
    
    @staticmethod
    def create_handlers(mode: str, config: dict) -> dict:
        """
        Create all mode-specific handlers.
        
        Args:
            mode: Mode string ('local' or 'googlecloud')
            config: Configuration dictionary (normalized)
            
        Returns:
            Dictionary containing all handlers:
            - auth: AuthenticationStrategy
            - image_source: ImageSourceStrategy
            - ai_client: AIClientStrategy
            - output: OutputStrategy
            - drive_service: Drive service (googlecloud) or None (local)
            - docs_service: Docs service (googlecloud) or None (local)
            
        Raises:
            ValueError: If mode is unknown
        """
        if mode == 'local':
            return ModeFactory._create_local_handlers(config)
        elif mode == 'googlecloud':
            return ModeFactory._create_googlecloud_handlers(config)
        else:
            raise ValueError(f"Unknown mode: {mode}")
    
    @staticmethod
    def _create_local_handlers(config: dict) -> dict:
        """
        Create handlers for local mode.
        
        Args:
            config: Configuration dictionary (normalized, with 'local' section)
            
        Returns:
            Dictionary containing all handlers for local mode
        """
        local_config = config['local']
        
        # Create authentication strategy
        auth = LocalAuthStrategy(local_config.get('api_key'))
        api_key = auth.authenticate()
        
        # Create image source strategy
        image_source = LocalImageSource(local_config['image_dir'])
        
        # Create output strategy (need ai_logger first for AI client)
        output_dir = local_config.get('output_dir', 'logs')
        ai_logger = logging.getLogger('ai_responses')
        image_dir = local_config['image_dir']
        
        # Create AI client strategy (pass ai_logger for response logging)
        model_id = local_config.get('ocr_model_id', 'gemini-3-flash-preview')
        ai_client = GeminiDevClient(api_key, model_id, ai_logger)
        
        # Create multiple output strategies (log, markdown, word)
        log_output = LogFileOutput(output_dir, ai_logger)
        md_output = MarkdownOutput(image_dir, config, None)  # prompt_text will be passed in initialize
        
        strategies = [log_output, md_output]
        output_formats = [f"Log ({output_dir}/)", f"Markdown ({image_dir}/)"]
        
        # Add Word output only if python-docx is available
        if DOCX_AVAILABLE:
            word_output = WordOutput(image_dir, config, None)  # prompt_text will be passed in initialize
            strategies.append(word_output)
            output_formats.append(f"Word ({image_dir}/)")
        else:
            logging.warning("python-docx not installed. Word output disabled. Install with: pip install python-docx>=0.8.11")
        
        # Wrap all strategies in CompositeOutput
        output = CompositeOutput(strategies)
        
        logging.info(f"Created LOCAL mode handlers: image_dir={image_dir}, output_dir={output_dir}, model={model_id}")
        logging.info(f"Output formats: {', '.join(output_formats)}")
        
        return {
            'auth': auth,
            'image_source': image_source,
            'ai_client': ai_client,
            'output': output,
            'drive_service': None,
            'docs_service': None
        }
    
    @staticmethod
    def _create_googlecloud_handlers(config: dict) -> dict:
        """
        Create handlers for Google Cloud mode.
        
        Args:
            config: Configuration dictionary (normalized, with 'googlecloud' section)
            
        Returns:
            Dictionary containing all handlers for Google Cloud mode
        """
        googlecloud_config = config['googlecloud']
        
        # Create authentication strategy and authenticate
        auth = GoogleCloudAuthStrategy(googlecloud_config['adc_file'])
        creds = auth.authenticate()
        
        # Initialize Google Cloud services
        drive_service, docs_service, genai_client = init_services(creds, googlecloud_config)
        
        # Create image source strategy
        image_source = DriveImageSource(
            drive_service,
            googlecloud_config['drive_folder_id'],
            document_name=googlecloud_config.get('document_name', 'Unknown')
        )
        
        # Create AI client strategy
        model_id = googlecloud_config.get('ocr_model_id', 'gemini-3-flash-preview')
        ai_client = VertexAIClient(genai_client, model_id)
        
        # Load prompt text for output strategy
        prompt_file = config.get('prompt_file', 'prompt.txt')
        try:
            with open(prompt_file, 'r', encoding='utf-8') as f:
                prompt_text = f.read()
        except Exception as e:
            logging.error(f"Error reading prompt file '{prompt_file}': {str(e)}")
            prompt_text = ""  # Fallback to empty prompt
        
        # Create output strategy
        # Pass full config (not just googlecloud_config) to include shared fields like archive_index
        output = GoogleDocsOutput(
            docs_service,
            drive_service,
            genai_client,
            config,  # Pass full normalized config, not just googlecloud section
            prompt_text
        )
        
        logging.info(f"Created GOOGLECLOUD mode handlers: project_id={googlecloud_config.get('project_id')}, drive_folder={googlecloud_config['drive_folder_id']}, model={model_id}")
        
        return {
            'auth': auth,
            'image_source': image_source,
            'ai_client': ai_client,
            'output': output,
            'drive_service': drive_service,
            'docs_service': docs_service,
            'genai_client': genai_client
        }


# ------------------------- EXISTING AUTHENTICATION & SERVICES -------------------------

def authenticate(adc_file: str):
    """
    Load user credentials from ADC file with required OAuth2 scopes.
    
    Args:
        adc_file: Path to the ADC credentials file
        
    Returns:
        Credentials object
        
    Raises:
        SystemExit: If token is expired or revoked, with helpful error message
    """
    scopes = [
        "https://www.googleapis.com/auth/drive",
        "https://www.googleapis.com/auth/documents",
        "https://www.googleapis.com/auth/cloud-platform"
    ]
    try:
        creds = Credentials.from_authorized_user_file(adc_file, scopes=scopes)
        
        # Check if credentials are expired and try to refresh if possible
        if creds.expired and creds.refresh_token:
            try:
                from google.auth.transport.requests import Request
                creds.refresh(Request())
                logging.info("Credentials refreshed successfully")
            except Exception as refresh_error:
                error_str = str(refresh_error).lower()
                if 'invalid_grant' in error_str or 'expired' in error_str or 'revoked' in error_str:
                    print("\n" + "="*70)
                    print("ERROR: Google Cloud authentication token has expired or been revoked")
                    print("="*70)
                    print("\nTo fix this, please refresh your Google Cloud credentials:")
                    print("\n  1. Verify your OAuth client configuration:")
                    print("     - Ensure client_secret.json belongs to the")
                    print("       correct Google Cloud project and Google account")
                    print("     - Check that the file matches the account you're authenticating with")
                    print("\n  2. Run the credential refresh script:")
                    print("     python refresh_credentials.py")
                    print("\n  3. Or use gcloud (if you have gcloud CLI installed):")
                    print("     gcloud auth application-default login --project=<PROJECT_ID> \\")
                    print("       --scopes=https://www.googleapis.com/auth/drive,https://www.googleapis.com/auth/documents,https://www.googleapis.com/auth/cloud-platform")
                    print("\n  4. After refreshing, run the transcription script again:")
                    print(f"     python transcribe.py <your_config.yaml>")
                    print("\n  For more details, see the 'Troubleshooting' section in README.md")
                    print("\n" + "="*70 + "\n")
                    logging.error(f"Google Cloud authentication failed: {str(refresh_error)}")
                    raise SystemExit(1)
                else:
                    raise
        
        logging.info(f"Credentials loaded with scopes: {scopes}")
        return creds
    except SystemExit:
        raise
    except Exception as e:
        error_str = str(e).lower()
        if 'invalid_grant' in error_str or 'expired' in error_str or 'revoked' in error_str:
            print("\n" + "="*70)
            print("ERROR: Google Cloud authentication token has expired or been revoked")
            print("="*70)
            print("\nTo fix this, please refresh your Google Cloud credentials:")
            print("\n  1. Verify your OAuth client configuration:")
            print("     - Ensure client_secret.json belongs to the")
            print("       correct Google Cloud project and Google account")
            print("     - Check that the file matches the account you're authenticating with")
            print("\n  2. Run the credential refresh script:")
            print("     python refresh_credentials.py")
            print("\n  3. Or use gcloud (if you have gcloud CLI installed):")
            print("     gcloud auth application-default login --project=<PROJECT_ID> \\")
            print("       --scopes=https://www.googleapis.com/auth/drive,https://www.googleapis.com/auth/documents,https://www.googleapis.com/auth/cloud-platform")
            print("\n  4. After refreshing, run the transcription script again:")
            print(f"     python transcribe.py <your_config.yaml>")
            print("\n  For more details, see the 'Troubleshooting' section in README.md")
            print("\n" + "="*70 + "\n")
            logging.error(f"Google Cloud authentication failed: {str(e)}")
            raise SystemExit(1)
        else:
            # Re-raise other authentication errors
            logging.error(f"Authentication error: {str(e)}")
            raise


def init_services(creds, config: dict):
    """
    Initialize Google Cloud services.
    
    Args:
        creds: OAuth2 credentials
        config: Configuration dictionary
    """
    project_id = config['project_id']
    region = config['region']
    adc_file = config['adc_file']
    
    logging.info(f"Initializing Vertex AI for project {project_id}...")
    # Set env var for Vertex AI SDK
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = adc_file
    
    # Initialize Gemini client
    genai_client = genai.Client(
        vertexai=True,
        project=project_id,
        location=region,
        credentials=creds
    )
    logging.info(f"Gemini client initialized in {region} with project {project_id}")

    logging.info("Initializing Google Drive and Docs APIs...")
    import httplib2
    from google_auth_httplib2 import AuthorizedHttp
    # Configure httplib2 with longer timeout for Google Docs API (5 minutes for large documents)
    http_base = httplib2.Http(timeout=300)  # 5 minutes timeout
    # Create authorized http object from credentials
    http = AuthorizedHttp(creds, http=http_base)
    drive = build("drive", "v3", http=http)
    docs = build("docs", "v1", http=http)
    logging.info("Google Drive and Docs APIs initialized with 5-minute timeout.")
    return drive, docs, genai_client


def extract_image_number(filename):
    """
    Extract the numeric identifier from an image filename.
    Supports the same patterns as list_images().
    Returns the extracted number, or None if no number can be extracted.
    """
    import re
    
    # Helper: case-insensitive check for JPEG extension
    def has_jpeg_extension(fname: str) -> bool:
        lower = fname.lower()
        return lower.endswith('.jpg') or lower.endswith('.jpeg')
    
    # Regex pattern for timestamp format: image - YYYY-MM-DDTHHMMSS.mmm.jpg/jpeg
    timestamp_pattern = re.compile(r'^image - (\d{4}-\d{2}-\d{2}T\d{6}\.\d{3})\.(?:jpg|jpeg)$', re.IGNORECASE)
    
    # Regex pattern for IMG_YYYYMMDD_XXXX.jpg format (e.g., IMG_20250814_0036.jpg)
    img_date_pattern = re.compile(r'^IMG_\d{8}_(\d+)\.(?:jpg|jpeg)$', re.IGNORECASE)
    
    number = None
    
    # Check for timestamp pattern first
    timestamp_match = timestamp_pattern.match(filename)
    if timestamp_match:
        # For timestamp images, we can't extract a meaningful number
        return None
    
    # Check for IMG_YYYYMMDD_XXXX.jpg pattern
    img_date_match = img_date_pattern.match(filename)
    if img_date_match:
        try:
            return int(img_date_match.group(1))
        except ValueError:
            pass
    
    # Check if filename matches the pattern image (N).jpg/jpeg
    if filename.startswith('image (') and (filename.lower().endswith(').jpg') or filename.lower().endswith(').jpeg')):
        try:
            start_idx = filename.find('(') + 1
            end_idx = filename.find(')')
            number_str = filename[start_idx:end_idx]
            return int(number_str)
        except (ValueError, IndexError):
            pass
    
    # Check if filename matches the pattern imageXXXXX.jpg/jpeg
    if filename.startswith('image') and has_jpeg_extension(filename) and '(' not in filename and ' - ' not in filename and '_' not in filename:
        try:
            ext_len = 5 if filename.lower().endswith('.jpeg') else 4
            number_str = filename[5:-ext_len]
            return int(number_str)
        except ValueError:
            pass
    
    # Check if filename matches the pattern XXXXX.jpg/jpeg
    if has_jpeg_extension(filename) and not filename.startswith('image') and '_' not in filename:
        try:
            ext_len = 5 if filename.lower().endswith('.jpeg') else 4
            number_str = filename[:-ext_len]
            return int(number_str)
        except ValueError:
            pass
    
    # Check if filename matches the pattern PREFIX_XXXXX.jpg/jpeg (e.g., 004933159_00216.jpeg)
    if has_jpeg_extension(filename) and '_' in filename:
        try:
            ext_len = 5 if filename.lower().endswith('.jpeg') else 4
            base_no_ext = filename[:-ext_len]
            # Take numeric part after the last underscore
            underscore_idx = base_no_ext.rfind('_')
            if underscore_idx != -1:
                suffix = base_no_ext[underscore_idx + 1:]
                if suffix.isdigit():
                    return int(suffix)
        except Exception:
            pass
    
    return None


def get_folder_name(drive_service, drive_folder_id: str):
    """
    Fetch folder name from Google Drive API using folder ID.
    
    Args:
        drive_service: Google Drive API service
        drive_folder_id: Google Drive folder ID
        
    Returns:
        Folder name string, or None if fetch fails
    """
    try:
        folder_metadata = drive_service.files().get(
            fileId=drive_folder_id,
            fields='name'
        ).execute()
        return folder_metadata.get('name')
    except Exception as e:
        logging.warning(f"Could not fetch folder name from Drive API: {str(e)}")
        return None


def list_images(drive_service, config: dict):
    """
    Get list of images from Google Drive folder, sorted by filename.
    Handles pagination to fetch up to max_images and filters based on image_start_number and image_count.
    Supports filename patterns: 
    - image (N).jpg where N is a number (e.g., image (7).jpg, image (10).jpg)
    - imageXXXXX.jpg where XXXXX is a 5-digit number (e.g., image00101.jpg)
    - XXXXX.jpg where XXXXX is a number (e.g., 52.jpg, 102.jpg)
    - XXXXXX_YYYYY.jpg where the number after underscore is used (e.g., 004933159_00216.jpeg)
    - IMG_YYYYMMDD_XXXX.jpg where XXXX is the image number (e.g., IMG_20250814_0036.jpg)
    - image - YYYY-MM-DDTHHMMSS.mmm.jpg (timestamp format, e.g., image - 2025-07-20T112914.366.jpg)
    Returns list of image metadata including id, name, and webViewLink.
    """
    import re
    from datetime import datetime
    
    # Handle both normalized (nested) and legacy (flat) config formats
    googlecloud_config = config.get('googlecloud', {})
    
    # Try nested format first, fall back to flat format for backward compatibility
    drive_folder_id = googlecloud_config.get('drive_folder_id') or config.get('drive_folder_id')
    if not drive_folder_id:
        raise KeyError("'drive_folder_id' not found in config (checked googlecloud.drive_folder_id and top-level)")
    
    document_name = googlecloud_config.get('document_name') or config.get('document_name', 'Unknown')
    max_images = config.get('max_images')
    if max_images is None:
        raise KeyError("'max_images' not found in config")
    retry_mode = config.get('retry_mode', False)
    retry_image_list = config.get('retry_image_list', [])
    image_start_number = config.get('image_start_number', 1)
    image_count = config.get('image_count')
    if image_count is None:
        raise KeyError("'image_count' not found in config")
    
    query = (
        f"mimeType='image/jpeg' and '{drive_folder_id}' in parents and trashed=false"
    )
    
    all_images = []
    page_token = None
    
    # Fetch all images with pagination (up to max_images)
    while len(all_images) < max_images:
        try:
            # Add orderBy parameter to sort by name
            resp = drive_service.files().list(
                q=query,
                fields="nextPageToken,files(id,name,webViewLink)",
                orderBy="name",  # Sort by filename
                pageSize=100,  # Fetch 100 files per request
                pageToken=page_token
            ).execute()
            
            files = resp.get('files', [])
            if not files:
                break
                
            all_images.extend(files)
            page_token = resp.get('nextPageToken')
            
            if not page_token:
                break
                
        except Exception as e:
            error_str = str(e).lower()
            if 'invalid_grant' in error_str or 'expired' in error_str or 'revoked' in error_str:
                logging.error(f"Error fetching images from Google Drive: {str(e)}")
                project_id = config.get('project_id', '<PROJECT_ID>')
                print("\n" + "="*70)
                print("ERROR: Google Cloud authentication token has expired or been revoked")
                print("="*70)
                print("\nTo fix this, please refresh your Google Cloud credentials:")
                print("\n  1. Verify your OAuth client configuration:")
                print("     - Ensure client_secret.json belongs to the")
                print("       correct Google Cloud project and Google account")
                print("     - Check that the file matches the account you're authenticating with")
                print("\n  2. Run the credential refresh script:")
                print("     python refresh_credentials.py")
                print("\n  3. Or use gcloud (if you have gcloud CLI installed):")
                print(f"     gcloud auth application-default login --project={project_id} \\")
                print("       --scopes=https://www.googleapis.com/auth/drive,https://www.googleapis.com/auth/documents,https://www.googleapis.com/auth/cloud-platform")
                print("\n  4. After refreshing, run the transcription script again:")
                print(f"     python transcribe.py <your_config.yaml>")
                print("\n  For more details, see the 'Troubleshooting' section in README.md")
                print("\n" + "="*70 + "\n")
                # Return empty list so the script can exit gracefully
                return []
            else:
                logging.error(f"Error fetching images from Google Drive: {str(e)}")
                break
    
    # Limit to max_images
    all_images = all_images[:max_images]
    logging.info(f"Found {len(all_images)} total images in folder (sorted by filename)")
    
    # RETRY MODE: If enabled, filter for specific failed images only
    if retry_mode:
        logging.info(f"RETRY MODE ENABLED: Looking for {len(retry_image_list)} specific failed images")
        retry_images = []
        
        # Convert retry list to full image names (add "image - " prefix if needed)
        retry_full_names = []
        for retry_img in retry_image_list:
            if retry_img.startswith('image - '):
                retry_full_names.append(retry_img)
            else:
                retry_full_names.append(f"image - {retry_img}")
        
        # Find matching images
        for img in all_images:
            if img['name'] in retry_full_names:
                retry_images.append(img)
        
        logging.info(f"Found {len(retry_images)} retry images out of {len(retry_image_list)} requested")
        if retry_images:
            retry_filenames = [img['name'] for img in retry_images]
            logging.info(f"Retry images found: {retry_filenames}")
        else:
            logging.warning("No retry images found! Check the retry_image_list names in config.")
        
        return retry_images
    
    # NORMAL MODE: Separate images by pattern type
    numbered_images = []
    timestamp_images = []
    
    # Helper: case-insensitive check for JPEG extension
    def has_jpeg_extension(filename: str) -> bool:
        lower = filename.lower()
        return lower.endswith('.jpg') or lower.endswith('.jpeg')

    # Regex pattern for timestamp format: image - YYYY-MM-DDTHHMMSS.mmm.jpg/jpeg
    timestamp_pattern = re.compile(r'^image - (\d{4}-\d{2}-\d{2}T\d{6}\.\d{3})\.(?:jpg|jpeg)$', re.IGNORECASE)
    
    # Regex pattern for IMG_YYYYMMDD_XXXX.jpg format (e.g., IMG_20250814_0036.jpg)
    img_date_pattern = re.compile(r'^IMG_\d{8}_(\d+)\.(?:jpg|jpeg)$', re.IGNORECASE)
    
    for img in all_images:
        filename = img['name']
        number = None
        timestamp_match = None
        img_date_match = None
        
        # Check for timestamp pattern first
        timestamp_match = timestamp_pattern.match(filename)
        if timestamp_match:
            timestamp_images.append(img)
            continue
        
        # Check for IMG_YYYYMMDD_XXXX.jpg pattern (e.g., IMG_20250814_0036.jpg)
        img_date_match = img_date_pattern.match(filename)
        if img_date_match:
            try:
                number = int(img_date_match.group(1))
            except ValueError:
                continue
        
        # Check if filename matches the pattern image (N).jpg/jpeg
        if filename.startswith('image (') and (filename.lower().endswith(').jpg') or filename.lower().endswith(').jpeg')):
            try:
                # Extract the number from filename (e.g., "image (7).jpg" -> 7)
                start_idx = filename.find('(') + 1
                end_idx = filename.find(')')
                number_str = filename[start_idx:end_idx]
                number = int(number_str)
            except (ValueError, IndexError):
                continue
        
        # Check if filename matches the pattern imageXXXXX.jpg/jpeg
        elif filename.startswith('image') and has_jpeg_extension(filename) and '(' not in filename and ' - ' not in filename and '_' not in filename:
            try:
                # Extract the number from filename (e.g., "image00101.jpg" -> 101)
                ext_len = 5 if filename.lower().endswith('.jpeg') else 4
                number_str = filename[5:-ext_len]  # Remove "image" prefix and extension suffix
                number = int(number_str)
            except ValueError:
                continue
        
        # Check if filename matches the pattern XXXXX.jpg/jpeg (like 52.jpg, 102.jpg)
        elif has_jpeg_extension(filename) and not filename.startswith('image') and '_' not in filename:
            try:
                # Extract the number from filename (e.g., "52.jpg" -> 52, "102.jpg" -> 102)
                ext_len = 5 if filename.lower().endswith('.jpeg') else 4
                number_str = filename[:-ext_len]  # Remove extension suffix
                number = int(number_str)
            except ValueError:
                continue

        # Check if filename matches the pattern PREFIX_XXXXX.jpg/jpeg (e.g., 004933159_00216.jpeg)
        elif has_jpeg_extension(filename) and '_' in filename:
            try:
                ext_len = 5 if filename.lower().endswith('.jpeg') else 4
                base_no_ext = filename[:-ext_len]
                # Take numeric part after the last underscore
                underscore_idx = base_no_ext.rfind('_')
                if underscore_idx != -1:
                    suffix = base_no_ext[underscore_idx + 1:]
                    if suffix.isdigit():
                        number = int(suffix)
            except Exception:
                continue
        
        # If we found a valid number, check if it's in the desired range
        if number is not None:
            if image_start_number <= number < image_start_number + image_count:
                numbered_images.append(img)
    
    # Handle numbered images (existing logic)
    filtered_images = []
    
    if numbered_images:
        # Define expected filename patterns for logging
        start_filename_pattern1 = f"image ({image_start_number}).jpg"
        end_filename_pattern1 = f"image ({image_start_number + image_count - 1}).jpg"
        start_filename_pattern2 = f"image{image_start_number:05d}.jpg"
        end_filename_pattern2 = f"image{image_start_number + image_count - 1:05d}.jpg"
        start_filename_pattern3 = f"{image_start_number}.jpg"
        end_filename_pattern3 = f"{image_start_number + image_count - 1}.jpg"
        
        logging.info(f"Filtering numbered images from {start_filename_pattern1} to {end_filename_pattern1} OR {start_filename_pattern2} to {end_filename_pattern2} OR {start_filename_pattern3} to {end_filename_pattern3}")
        
        # Sort numbered images numerically by their extracted number
        def extract_number_for_sorting(img):
            filename = img['name']
            lower = filename.lower()
            if filename.startswith('image (') and (lower.endswith(').jpg') or lower.endswith(').jpeg')):
                # Extract number from "image (N).jpg" format
                start_idx = filename.find('(') + 1
                end_idx = filename.find(')')
                number_str = filename[start_idx:end_idx]
            elif filename.startswith('image') and has_jpeg_extension(filename) and '(' not in filename and ' - ' not in filename and '_' not in filename:
                # Extract number from "imageXXXXX.jpg" format
                ext_len = 5 if lower.endswith('.jpeg') else 4
                number_str = filename[5:-ext_len]  # Remove "image" prefix and extension suffix
            elif has_jpeg_extension(filename) and '_' in filename:
                # Extract number from "PREFIX_XXXXX.jpg/jpeg" format
                ext_len = 5 if lower.endswith('.jpeg') else 4
                base_no_ext = filename[:-ext_len]
                underscore_idx = base_no_ext.rfind('_')
                number_str = base_no_ext[underscore_idx + 1:]
            else:
                # Extract number from "XXXXX.jpg" format
                ext_len = 5 if lower.endswith('.jpeg') else 4
                number_str = filename[:-ext_len]  # Remove extension suffix
            return int(number_str)
        
        numbered_images.sort(key=extract_number_for_sorting)
        filtered_images.extend(numbered_images)
    
    # Handle timestamp images
    if timestamp_images:
        logging.info(f"Found {len(timestamp_images)} timestamp-based images")
        
        # Sort timestamp images chronologically
        def extract_timestamp_for_sorting(img):
            filename = img['name']
            match = timestamp_pattern.match(filename)
            if match:
                timestamp_str = match.group(1)
                # Parse timestamp: YYYY-MM-DDTHHMMSS.mmm -> datetime
                try:
                    # Convert format 2025-07-20T112914.366 to 2025-07-20T11:29:14.366
                    formatted_timestamp = f"{timestamp_str[:11]}{timestamp_str[11:13]}:{timestamp_str[13:15]}:{timestamp_str[15:]}"
                    return datetime.fromisoformat(formatted_timestamp)
                except ValueError:
                    return datetime.min
            return datetime.min
        
        timestamp_images.sort(key=extract_timestamp_for_sorting)
        
        # For timestamp images, treat image_start_number as the starting position (1-indexed)
        # and image_count as the number of images to process
        start_pos = max(1, image_start_number) - 1  # Convert to 0-indexed
        end_pos = min(len(timestamp_images), start_pos + image_count)
        
        selected_timestamp_images = timestamp_images[start_pos:end_pos]
        filtered_images.extend(selected_timestamp_images)
        
        if selected_timestamp_images:
            logging.info(f"Selected {len(selected_timestamp_images)} timestamp images from position {image_start_number} to {start_pos + len(selected_timestamp_images)}")
            filenames = [img['name'] for img in selected_timestamp_images]
            logging.info(f"Selected timestamp files: {filenames}")
    
    # Final sort of all filtered images
    if filtered_images:
        # Sort mixed list: numbered images by number, timestamp images by timestamp
        def mixed_sorting_key(img):
            filename = img['name']
            timestamp_match = timestamp_pattern.match(filename)
            
            if timestamp_match:
                # Timestamp image - sort by timestamp
                timestamp_str = timestamp_match.group(1)
                try:
                    formatted_timestamp = f"{timestamp_str[:11]}{timestamp_str[11:13]}:{timestamp_str[13:15]}:{timestamp_str[15:]}"
                    return (1, datetime.fromisoformat(formatted_timestamp))  # 1 to sort timestamp images after numbered
                except ValueError:
                    return (1, datetime.min)
            else:
                # Numbered image - sort by number
                lower = filename.lower()
                if filename.startswith('image (') and (lower.endswith(').jpg') or lower.endswith(').jpeg')):
                    start_idx = filename.find('(') + 1
                    end_idx = filename.find(')')
                    number_str = filename[start_idx:end_idx]
                elif filename.startswith('image') and has_jpeg_extension(filename) and '(' not in filename and ' - ' not in filename and '_' not in filename:
                    ext_len = 5 if lower.endswith('.jpeg') else 4
                    number_str = filename[5:-ext_len]
                elif has_jpeg_extension(filename) and '_' in filename:
                    ext_len = 5 if lower.endswith('.jpeg') else 4
                    base_no_ext = filename[:-ext_len]
                    underscore_idx = base_no_ext.rfind('_')
                    number_str = base_no_ext[underscore_idx + 1:]
                else:
                    ext_len = 5 if lower.endswith('.jpeg') else 4
                    number_str = filename[:-ext_len]
                try:
                    return (0, int(number_str))  # 0 to sort numbered images first
                except ValueError:
                    return (0, 0)
        
        filtered_images.sort(key=mixed_sorting_key)
    
    # If no images were selected by number/timestamp filters, fall back to position-based selection
    if not filtered_images and all_images and not retry_mode:
        # all_images are already sorted by name from the Drive API
        start_pos = max(1, image_start_number) - 1
        end_pos = min(len(all_images), start_pos + image_count)
        fallback_selected = all_images[start_pos:end_pos]
        if fallback_selected:
            logging.info(f"No numeric/timestamp matches; falling back to position selection: items {image_start_number} to {image_start_number + len(fallback_selected) - 1}")
            filtered_images = fallback_selected

    logging.info(f"Selected {len(filtered_images)} total images for processing")
    
    # Log the selected filenames for verification
    if filtered_images:
        filenames = [img['name'] for img in filtered_images]
        logging.info(f"Final selected files: {filenames}")
    
    return filtered_images


def download_image(drive_service, file_id, file_name, document_name: str):
    import time
    download_start = time.time()
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Downloading image '{file_name}'")
    try:
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        chunk_count = 0
        while not done:
            status, done = downloader.next_chunk()
            chunk_count += 1
            if status:
                progress = int(status.progress() * 100)
                logging.debug(f"[{datetime.now().strftime('%H:%M:%S')}] Download progress for '{file_name}': {progress}%")
        fh.seek(0)
        img_bytes = fh.read()
        download_elapsed = time.time() - download_start
        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Image '{file_name}' downloaded successfully ({len(img_bytes)} bytes) in {download_elapsed:.1f}s ({chunk_count} chunks)")
        
        if download_elapsed > 30:
            logging.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: Download took {download_elapsed:.1f}s (>30s) for '{file_name}' - possible network issues")
        
        return img_bytes
    except Exception as e:
        download_elapsed = time.time() - download_start
        error_type = type(e).__name__
        logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] Error downloading '{file_name}' after {download_elapsed:.1f}s: {error_type}: {str(e)}")
        logging.error(f"Full traceback:\n{traceback.format_exc()}")
        raise


def transcribe_image(genai_client, image_bytes, file_name, prompt_text: str, ocr_model_id: str):
    import signal
    import time
    
    function_start_time = time.time()
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Starting transcription for image '{file_name}' (size: {len(image_bytes)} bytes)")
    ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === Starting transcription for {file_name} ===")
    
    # Create image part using base64 encoding
    image_part = types.Part.from_bytes(
        data=image_bytes,
        mime_type="image/jpeg"
    )
    
            # Create content with instruction and image
    content = types.Content(
        role="user",
        parts=[
                    types.Part.from_text(text=prompt_text),
            image_part
        ]
    )
    
    # Configure generation parameters
    generate_content_config = types.GenerateContentConfig(
        temperature=0.1,
        top_p=0.8,
        seed=0,
        max_output_tokens=65535,
        # safety_settings=[
        #     types.SafetySetting(
        #         category="HARM_CATEGORY_HATE_SPEECH",
        #         threshold="OFF"
        #     ),
        #     types.SafetySetting(
        #         category="HARM_CATEGORY_DANGEROUS_CONTENT",
        #         threshold="OFF"
        #     ),
        #     types.SafetySetting(
        #         category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
        #         threshold="OFF"
        #     ),
        #     types.SafetySetting(
        #         category="HARM_CATEGORY_HARASSMENT",
        #         threshold="OFF"
        #     )
        # ],
        system_instruction=[types.Part.from_text(text=prompt_text)],
        thinking_config=types.ThinkingConfig(
            thinking_budget=5000,
        ),
    )
    
    max_retries = 3
    retry_delay = 30  # seconds
    # Exponential backoff timeouts: 1 min, 2 min, 5 min
    timeout_seconds_list = [60, 120, 300]
    
    for attempt in range(max_retries):
        attempt_start_time = time.time()
        timeout_seconds = timeout_seconds_list[attempt]
        
        # Define timeout handler inside loop to properly capture timeout_seconds
        def timeout_handler(signum, frame):
            elapsed = time.time() - function_start_time
            error_msg = f"Vertex AI API call timed out after {timeout_seconds/60:.1f} minutes (total elapsed: {elapsed:.1f}s) for {file_name}"
            logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] {error_msg}")
            ai_logger.error(f"[{datetime.now().strftime('%H:%M:%S')}] TIMEOUT: {error_msg}")
            raise TimeoutError(error_msg)
        
        try:
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} for image '{file_name}' (timeout: {timeout_seconds/60:.1f} min)")
            ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} starting for {file_name} (timeout: {timeout_seconds/60:.1f} min)")
            
            # Set up timeout with exponential backoff (1 min, 2 min, 5 min)
            signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(timeout_seconds)
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Timeout set to {timeout_seconds/60:.1f} minutes for '{file_name}' (attempt {attempt + 1}/{max_retries})")
            
            api_call_start = time.time()
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Making API call to Vertex AI for '{file_name}'...")
            ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] API call initiated for {file_name}")
            
            # Generate content
            response = genai_client.models.generate_content(
                model=ocr_model_id,
                contents=[content],
                config=generate_content_config
            )
            
            # Cancel the timeout
            signal.alarm(0)
            
            api_call_elapsed = time.time() - api_call_start
            elapsed_time = time.time() - attempt_start_time
            total_elapsed = time.time() - function_start_time
            
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Vertex AI response received in {api_call_elapsed:.1f} seconds (attempt total: {elapsed_time:.1f}s, function total: {total_elapsed:.1f}s) for '{file_name}'")
            ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] API call completed in {api_call_elapsed:.1f}s for {file_name}")
            
            # Log warning if API call took unusually long
            if api_call_elapsed > 60:
                logging.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: API call took {api_call_elapsed:.1f} seconds (>60s) for '{file_name}' - possible throttling or network issues")
                ai_logger.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: Long API call duration ({api_call_elapsed:.1f}s) for {file_name}")
            
            text = response.text
            
            # Ensure text is not None
            if text is None:
                text = "[No response text received from Vertex AI]"
            
            # Log the full AI response to the AI responses log
            ai_logger.info(f"=== AI Response for {file_name} ===")
            ai_logger.info(f"Model: {ocr_model_id}")
            ai_logger.info(f"Request timestamp: {datetime.now().isoformat()}")
            ai_logger.info(f"Image size: {len(image_bytes)} bytes")
            ai_logger.info(f"Instruction length: {len(prompt_text)} characters")
            ai_logger.info(f"Response length: {len(text) if text else 0} characters")
            ai_logger.info(f"Processing time: {elapsed_time:.1f} seconds")
            
            # Log response metadata if available
            if hasattr(response, 'usage_metadata'):
                ai_logger.info(f"Usage metadata: {response.usage_metadata}")
            if hasattr(response, 'candidates') and response.candidates:
                ai_logger.info(f"Number of candidates: {len(response.candidates)}")
                if hasattr(response.candidates[0], 'finish_reason'):
                    ai_logger.info(f"Finish reason: {response.candidates[0].finish_reason}")
            
            ai_logger.info(f"Full response:\n{text}")
            ai_logger.info(f"=== End AI Response for {file_name} ===\n")
            
            # Log the first and last few lines of the response for troubleshooting
            if text:
                lines = text.split('\n')
                if len(lines) > 6:
                    first_lines = '\n'.join(lines[:3])
                    last_lines = '\n'.join(lines[-3:])
                    logging.info(f"First 3 lines of transcription for '{file_name}':\n{first_lines}\n...\nLast 3 lines:\n{last_lines}")
                else:
                    logging.info(f"Full transcription for '{file_name}':\n{text}")
            
            # Return text, timing, and usage metadata
            usage_metadata = None
            if hasattr(response, 'usage_metadata'):
                usage_metadata = response.usage_metadata
            
            function_total_elapsed = time.time() - function_start_time
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Transcription function completed for '{file_name}' in {function_total_elapsed:.1f}s total")
            ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Transcription completed successfully for {file_name} (total: {function_total_elapsed:.1f}s)")
            
            return text, elapsed_time, usage_metadata
            
        except (TimeoutError, ConnectionError, OSError) as e:
            # Cancel any pending timeout
            signal.alarm(0)
            
            attempt_elapsed = time.time() - attempt_start_time
            total_elapsed = time.time() - function_start_time
            error_type = type(e).__name__
            
            error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} failed for '{file_name}' after {attempt_elapsed:.1f}s (total elapsed: {total_elapsed:.1f}s): {error_type}: {str(e)}"
            logging.warning(error_msg)
            ai_logger.warning(f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1} failed for {file_name}: {error_type}: {str(e)}")
            ai_logger.warning(f"Attempt elapsed time: {attempt_elapsed:.1f}s, Total function time: {total_elapsed:.1f}s")
            
            # Log full traceback for debugging
            logging.debug(f"Full traceback for {file_name}:\n{traceback.format_exc()}")
            ai_logger.debug(f"Traceback for {file_name}:\n{traceback.format_exc()}")
            
            if attempt < max_retries - 1:
                next_timeout = timeout_seconds_list[attempt + 1]
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Retrying in {retry_delay} seconds... (exponential backoff, next timeout: {next_timeout/60:.1f} min)")
                ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] Will retry in {retry_delay}s with exponential backoff (next timeout: {next_timeout/60:.1f} min)")
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Retry delay completed, starting attempt {attempt + 2}/{max_retries}...")
                continue  # Explicitly continue to next iteration
            else:
                error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] All {max_retries} attempts failed for '{file_name}' after {total_elapsed:.1f}s total: {error_type}: {str(e)}"
                ai_logger.error(f"[{datetime.now().strftime('%H:%M:%S')}] === AI Error for {file_name} ===")
                ai_logger.error(f"Error type: {error_type}")
                ai_logger.error(f"Error message: {str(e)}")
                ai_logger.error(f"Total elapsed time: {total_elapsed:.1f}s")
                ai_logger.error(f"Attempts made: {max_retries}")
                ai_logger.error(f"Full traceback:\n{traceback.format_exc()}")
                ai_logger.error(f"=== End AI Error for {file_name} ===\n")
                logging.error(error_msg)
                logging.error(f"All {max_retries} retry attempts exhausted for '{file_name}'")
                logging.error(f"Full traceback:\n{traceback.format_exc()}")
                # Return error text with None for timing and metadata
                return f"[Error during transcription: {str(e)}]", None, None
        
        except Exception as e:
            # Check if this is a timeout-related exception from httpx/httpcore
            # httpx.ReadTimeout and httpcore.ReadTimeout should be retried
            error_type = type(e).__name__
            error_module = type(e).__module__
            error_str = str(e).lower()
            
            # Detect timeout errors from various sources
            is_timeout_error = (
                'Timeout' in error_type or  # TimeoutError, ReadTimeout, etc.
                'ReadTimeout' in error_type or  # httpx.ReadTimeout, httpcore.ReadTimeout
                ('timeout' in error_str and ('httpx' in error_module or 'httpcore' in error_module))  # Timeout errors from httpx/httpcore
            )
            
            # If it's a timeout error, treat it like TimeoutError and retry
            if is_timeout_error and attempt < max_retries - 1:
                # Cancel any pending timeout
                signal.alarm(0)
                
                attempt_elapsed = time.time() - attempt_start_time
                total_elapsed = time.time() - function_start_time
                
                error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} failed for '{file_name}' after {attempt_elapsed:.1f}s (total elapsed: {total_elapsed:.1f}s): {error_type}: {str(e)}"
                logging.warning(error_msg)
                ai_logger.warning(f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1} failed for {file_name}: {error_type}: {str(e)}")
                ai_logger.warning(f"Attempt elapsed time: {attempt_elapsed:.1f}s, Total function time: {total_elapsed:.1f}s")
                
                next_timeout = timeout_seconds_list[attempt + 1]
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Retrying in {retry_delay} seconds... (exponential backoff, next timeout: {next_timeout/60:.1f} min)")
                ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] Will retry in {retry_delay}s with exponential backoff (next timeout: {next_timeout/60:.1f} min)")
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Retry delay completed, starting attempt {attempt + 2}/{max_retries}...")
                continue  # Explicitly continue to next iteration
                
            # Not a timeout error or all retries exhausted - handle as unexpected error
            # Cancel any pending timeout
            signal.alarm(0)
            
            attempt_elapsed = time.time() - attempt_start_time
            total_elapsed = time.time() - function_start_time
            error_type = type(e).__name__
            
            error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Unexpected error in Vertex AI transcription for '{file_name}' after {attempt_elapsed:.1f}s (total elapsed: {total_elapsed:.1f}s): {error_type}: {str(e)}"
            ai_logger.error(f"[{datetime.now().strftime('%H:%M:%S')}] === AI Error for {file_name} ===")
            ai_logger.error(f"Error type: {error_type}")
            ai_logger.error(f"Error message: {str(e)}")
            ai_logger.error(f"Attempt elapsed time: {attempt_elapsed:.1f}s, Total function time: {total_elapsed:.1f}s")
            ai_logger.error(f"Full traceback:\n{traceback.format_exc()}")
            ai_logger.error(f"=== End AI Error for {file_name} ===\n")
            logging.error(error_msg)
            logging.error(f"Full traceback:\n{traceback.format_exc()}")
            # Return error text with None for timing and metadata
            return f"[Error during transcription: {str(e)}]", None, None


def find_title_page_image(drive_service, drive_folder_id: str, title_page_filename: str):
    """
    Find title page image file in Google Drive folder by filename.
    
    Args:
        drive_service: Google Drive API service
        drive_folder_id: Google Drive folder ID
        title_page_filename: Filename of the title page image
        
    Returns:
        File ID if found, None otherwise
    """
    try:
        query = (
            f"name='{title_page_filename}' and '{drive_folder_id}' in parents and "
            f"(mimeType='image/jpeg' or mimeType='image/jpg' or mimeType='image/png') and trashed=false"
        )
        results = drive_service.files().list(
            q=query,
            fields="files(id, name)"
        ).execute()
        
        files = results.get('files', [])
        if files:
            return files[0]['id']
        return None
    except Exception as e:
        logging.warning(f"Error searching for title page image '{title_page_filename}': {str(e)}")
        return None


def upload_image_to_drive(drive_service, image_bytes, filename: str, drive_folder_id: str):
    """
    Upload image bytes to Google Drive and return the file ID.
    
    Args:
        drive_service: Google Drive API service
        image_bytes: Image data as bytes
        filename: Filename for the uploaded image
        drive_folder_id: Google Drive folder ID where to upload
        
    Returns:
        File ID of uploaded image, or None if upload fails
    """
    try:
        from googleapiclient.http import MediaIoBaseUpload
        
        file_metadata = {
            'name': filename,
            'parents': [drive_folder_id]
        }
        media = MediaIoBaseUpload(io.BytesIO(image_bytes), mimetype='image/jpeg', resumable=True)
        
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        
        return file.get('id')
    except Exception as e:
        logging.error(f"Error uploading image to Drive: {str(e)}")
        return None


def upload_log_file_to_drive(drive_service, log_file_path: str, drive_folder_id: str) -> str:
    """
    Upload a log file to Google Drive and return the file ID.
    
    Args:
        drive_service: Google Drive API service
        log_file_path: Local path to the log file
        drive_folder_id: Google Drive folder ID where to upload
        
    Returns:
        File ID of uploaded log file, or None if upload fails
    """
    try:
        from googleapiclient.http import MediaFileUpload
        
        if not os.path.exists(log_file_path):
            logging.warning(f"Log file not found: {log_file_path}")
            return None
        
        filename = os.path.basename(log_file_path)
        
        file_metadata = {
            'name': filename,
            'parents': [drive_folder_id]
        }
        
        media = MediaFileUpload(log_file_path, mimetype='text/plain', resumable=True)
        
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name'
        ).execute()
        
        file_id = file.get('id')
        logging.info(f"Uploaded log file '{filename}' to Google Drive folder (ID: {file_id})")
        return file_id
    except Exception as e:
        logging.error(f"Error uploading log file to Drive: {str(e)}")
        return None


def insert_title_page_image_and_transcribe(docs_service, drive_service, doc_id: str, config: dict, insert_index: int, genai_client, prompt_text: str):
    """
    Insert title page image into Google Doc and transcribe it.
    If image insertion fails (e.g., sharing not allowed), still transcribe and return transcription.
    
    Args:
        docs_service: Google Docs API service
        drive_service: Google Drive API service
        doc_id: Google Doc ID
        config: Configuration dictionary
        insert_index: Index where to insert the image
        genai_client: Vertex AI Gemini client for transcription
        prompt_text: Prompt text for transcription
        
    Returns:
        Tuple of (new_index, transcription_text) where transcription_text is the transcribed content or None
    """
    title_page_filename = config.get('title_page_filename')
    if not title_page_filename:
        return insert_index, None
    
    drive_folder_id = config['drive_folder_id']
    ocr_model_id = config.get('ocr_model_id', 'gemini-3-flash-preview')
    transcription_text = None
    image_inserted = False
    
    try:
        # Try to find the image in the Drive folder
        image_file_id = find_title_page_image(drive_service, drive_folder_id, title_page_filename)
        
        if not image_file_id:
            logging.warning(f"Title page image '{title_page_filename}' not found in Drive folder")
            return insert_index, None
        
        # Try to make the file publicly accessible temporarily (required for Docs API)
        file_public = False
        try:
            drive_service.permissions().create(
                fileId=image_file_id,
                body={'role': 'reader', 'type': 'anyone'},
                fields='id'
            ).execute()
            file_public = True
            logging.info(f"Made title page image '{title_page_filename}' publicly accessible for insertion")
        except Exception as perm_error:
            # Check if permission already exists
            if 'duplicate' in str(perm_error).lower() or 'already' in str(perm_error).lower():
                file_public = True
                logging.info(f"Title page image '{title_page_filename}' is already publicly accessible")
            else:
                logging.warning(f"Could not make title page image publicly accessible (sharing may not be allowed): {str(perm_error)}")
                logging.info(f"Will skip image insertion but will still transcribe the image")
        
        # Try to insert the image if we made it public (or it was already public)
        if file_public:
            try:
                # Get webContentLink for the image
                file_metadata = drive_service.files().get(
                    fileId=image_file_id,
                    fields='id, name, webContentLink'
                ).execute()
                
                image_uri = file_metadata.get('webContentLink')
                if not image_uri:
                    # Fallback to direct view link
                    image_uri = f"https://drive.google.com/uc?export=view&id={image_file_id}"
                
                # Insert the image
                requests = [{
                    'insertInlineImage': {
                        'location': {'index': insert_index},
                        'uri': image_uri,
                        'objectSize': {
                            'height': {'magnitude': 400, 'unit': 'PT'},
                            'width': {'magnitude': 600, 'unit': 'PT'}
                        }
                    }
                }]
                
                docs_service.documents().batchUpdate(
                    documentId=doc_id,
                    body={'requests': requests}
                ).execute()
                
                logging.info(f"Title page image '{title_page_filename}' inserted successfully")
                image_inserted = True
                
                # Refresh index after image insertion and add a newline after the image
                doc = docs_service.documents().get(documentId=doc_id).execute()
                insert_index = doc['body']['content'][-1]['endIndex'] - 1
                
                # Insert a newline after the image to ensure proper spacing
                # This prevents the image from being affected by subsequent text insertions
                newline_request = [{
                    'insertText': {
                        'location': {'index': insert_index},
                        'text': '\n'
                    }
                }]
                docs_service.documents().batchUpdate(
                    documentId=doc_id,
                    body={'requests': newline_request}
                ).execute()
                
                # Refresh index again after newline insertion
                doc = docs_service.documents().get(documentId=doc_id).execute()
                insert_index = doc['body']['content'][-1]['endIndex'] - 1
                
            except Exception as img_error:
                logging.warning(f"Error inserting title page image (will still transcribe): {str(img_error)}")
                image_inserted = False
        
        # Always try to download and transcribe the title page image (even if insertion failed)
        try:
            logging.info(f"Transcribing title page image '{title_page_filename}'...")
            img_bytes = download_image(drive_service, image_file_id, title_page_filename, config.get('document_name', 'Unknown'))
            transcription_text, _, _ = transcribe_image(genai_client, img_bytes, title_page_filename, prompt_text, ocr_model_id)
            
            if transcription_text and not transcription_text.startswith('[Error'):
                logging.info(f"Title page image '{title_page_filename}' transcribed successfully")
            else:
                logging.warning(f"Title page image transcription failed or returned error")
                transcription_text = None
        except Exception as e:
            logging.warning(f"Error transcribing title page image: {str(e)}")
            transcription_text = None
        
        return insert_index, transcription_text
        
    except Exception as e:
        logging.warning(f"Error in title page image processing: {str(e)}")
        return insert_index, None


def create_doc(docs_service, drive_service, title, config: dict):
    """Create a new Google Doc in the specified folder and return its ID."""
    # Handle both normalized (nested) and legacy (flat) config formats
    googlecloud_config = config.get('googlecloud', {})
    drive_folder_id = googlecloud_config.get('drive_folder_id') or config.get('drive_folder_id')
    if not drive_folder_id:
        raise KeyError("'drive_folder_id' not found in config (checked googlecloud.drive_folder_id and top-level)")
    
    try:
        # First create the document
        doc = docs_service.documents().create(body={'title': title}).execute()
        doc_id = doc['documentId']
        
        # Then move it to the specified folder
        file = drive_service.files().update(
            fileId=doc_id,
            addParents=drive_folder_id,
            fields='id, parents'
        ).execute()
        
        logging.info(f"Created new Google Doc '{title}' with ID: {doc_id}")
        return doc_id
    except Exception as e:
        logging.error(f"Error creating Google Doc: {str(e)}")
        # Check if it's a permission error
        if 'insufficientFilePermissions' in str(e) or '403' in str(e):
            logging.warning(f"Insufficient permissions to add document to folder")
            logging.warning(f"Document was created but could not be moved to the target folder")
            logging.info(f"Returning None to trigger local save fallback")
            return None
        raise


def _log_run_summary(output, transcribed_pages, metrics, start_time, end_time, error_info, mode):
    """
    Log comprehensive Run Summary to main logger.
    
    Args:
        output: Output strategy object (may be CompositeOutput)
        transcribed_pages: List of transcribed pages
        metrics: Metrics dictionary
        start_time: Session start datetime
        end_time: Session end datetime
        error_info: Optional error info dict
        mode: 'local' or 'googlecloud'
    """
    logging.info("")
    logging.info("=" * 80)
    logging.info("RUN SUMMARY")
    logging.info("=" * 80)
    
    # Status
    if error_info:
        logging.info("Status: INTERRUPTED BY ERROR")
        logging.info(f"  Error Type: {error_info.get('type', 'Unknown')}")
        logging.info(f"  Error Message: {error_info.get('message', 'Unknown error')}")
        if error_info.get('next_image_number'):
            logging.info(f"  Resume: Update config 'image_start_number' to {error_info['next_image_number']}")
    else:
        logging.info("Status: COMPLETED SUCCESSFULLY")
    
    # Statistics
    total_images = len(transcribed_pages)
    successful = len([p for p in transcribed_pages if p.get('text') and not p['text'].startswith('[Error')])
    failed = total_images - successful
    
    logging.info("")
    logging.info("Statistics:")
    logging.info(f"  Total images processed: {total_images}")
    logging.info(f"  Successful transcriptions: {successful}")
    logging.info(f"  Failed transcriptions: {failed}")
    
    # Timing
    if start_time and end_time:
        elapsed = (end_time - start_time).total_seconds()
        logging.info(f"  Start time: {start_time.strftime('%Y-%m-%d %H:%M:%S')}")
        logging.info(f"  End time: {end_time.strftime('%Y-%m-%d %H:%M:%S')}")
        logging.info(f"  Total duration: {elapsed:.1f} seconds ({elapsed/60:.1f} minutes)")
    
    # Metrics
    if metrics and isinstance(metrics, dict) and len(metrics) > 0:
        logging.info("")
        logging.info("Metrics:")
        for key, value in metrics.items():
            if isinstance(value, (int, float)):
                if 'token' in key.lower() or 'count' in key.lower():
                    logging.info(f"  {key}: {value:,}")
                elif 'cost' in key.lower() or 'price' in key.lower():
                    logging.info(f"  {key}: ${value:.4f}")
                elif 'time' in key.lower() or 'duration' in key.lower() or 'elapsed' in key.lower():
                    logging.info(f"  {key}: {value:.2f}s")
                else:
                    logging.info(f"  {key}: {value}")
            else:
                logging.info(f"  {key}: {value}")
    elif metrics is None:
        logging.info("")
        logging.info("Metrics: (Not available)")
    
    # Outputs produced
    logging.info("")
    logging.info("Outputs Produced:")
    
    def _collect_outputs(strategy, outputs_list):
        """Recursively collect output paths from strategy (handles CompositeOutput)."""
        if isinstance(strategy, CompositeOutput):
            for sub_strategy in strategy.strategies:
                _collect_outputs(sub_strategy, outputs_list)
        else:
            # LogFileOutput
            if hasattr(strategy, 'log_file_path') and strategy.log_file_path:
                outputs_list.append(('Log File', strategy.log_file_path))
            # MarkdownOutput
            if hasattr(strategy, 'final_file_path') and strategy.final_file_path:
                outputs_list.append(('Markdown File', strategy.final_file_path))
            # WordOutput
            if hasattr(strategy, 'doc_path') and strategy.doc_path:
                outputs_list.append(('Word Document', strategy.doc_path))
            # GoogleDocsOutput
            if hasattr(strategy, 'doc_id') and strategy.doc_id:
                doc_url = f"https://docs.google.com/document/d/{strategy.doc_id}"
                outputs_list.append(('Google Doc', doc_url))
    
    outputs = []
    _collect_outputs(output, outputs)
    
    if outputs:
        for output_type, output_path in outputs:
            logging.info(f"  {output_type}: {output_path}")
    else:
        logging.info("  (No outputs generated)")
    
    logging.info("=" * 80)
    logging.info("")


def calculate_metrics(usage_metadata_list, timing_list):
    """
    Calculate metrics from usage metadata and timing data.
    Returns a dictionary with calculated metrics.
    """
    # Pricing rates (per 1M tokens) - estimated 2026 pricing for Gemini 3 Flash
    INPUT_RATE = 0.15  # $0.15 per 1M input tokens
    OUTPUT_RATE = 0.60  # $0.60 per 1M output tokens
    CACHED_RATE = 0.03  # $0.03 per 1M cached tokens
    
    total_time = sum(t for t in timing_list if t is not None)
    page_count = len([t for t in timing_list if t is not None])
    avg_time_per_page = total_time / page_count if page_count > 0 else 0
    
    total_input_tokens = 0
    total_output_tokens = 0
    total_cached_tokens = 0
    
    for usage_metadata in usage_metadata_list:
        if usage_metadata is None:
            continue
        
        # Handle both dict and object formats
        if isinstance(usage_metadata, dict):
            # Dictionary format (from LOCAL mode)
            prompt_tokens = usage_metadata.get('prompt_tokens', 0) or 0
            completion_tokens = usage_metadata.get('completion_tokens', 0) or 0
            cached_tokens = usage_metadata.get('cached_tokens', 0) or 0
            thoughts_tokens = usage_metadata.get('thoughts_tokens', 0) or 0
            
            total_input_tokens += prompt_tokens
            total_output_tokens += completion_tokens
            if thoughts_tokens:
                total_output_tokens += thoughts_tokens
            total_cached_tokens += cached_tokens
        else:
            # Object format (from GOOGLECLOUD mode)
            # Extract prompt tokens (input) - this includes both text and image tokens
            if hasattr(usage_metadata, 'prompt_token_count') and usage_metadata.prompt_token_count:
                total_input_tokens += usage_metadata.prompt_token_count
            
            # Extract cached tokens (billed at lower rate)
            if hasattr(usage_metadata, 'cached_content_token_count') and usage_metadata.cached_content_token_count:
                total_cached_tokens += usage_metadata.cached_content_token_count
            
            # Extract candidate tokens (output) - this is the visible response
            if hasattr(usage_metadata, 'candidates_token_count') and usage_metadata.candidates_token_count:
                total_output_tokens += usage_metadata.candidates_token_count
            
            # Extract thoughts tokens (billed as output) - reasoning tokens
            if hasattr(usage_metadata, 'thoughts_token_count') and usage_metadata.thoughts_token_count:
                total_output_tokens += usage_metadata.thoughts_token_count
    
    # Calculate costs
    # Input tokens (excluding cached, which are billed separately)
    input_tokens_billed = total_input_tokens - total_cached_tokens if total_cached_tokens > 0 else total_input_tokens
    input_cost = (input_tokens_billed / 1_000_000) * INPUT_RATE
    output_cost = (total_output_tokens / 1_000_000) * OUTPUT_RATE
    cached_cost = (total_cached_tokens / 1_000_000) * CACHED_RATE
    estimated_cost_per_run = input_cost + output_cost + cached_cost
    estimated_cost_per_page = estimated_cost_per_run / page_count if page_count > 0 else 0
    
    return {
        'total_time': total_time,
        'avg_time_per_page': avg_time_per_page,
        'total_input_tokens': total_input_tokens,
        'total_output_tokens': total_output_tokens,
        'total_cached_tokens': total_cached_tokens,
        'estimated_cost_per_run': estimated_cost_per_run,
        'estimated_cost_per_page': estimated_cost_per_page,
        'page_count': page_count
    }


def create_overview_section(pages, config: dict, prompt_text: str, metrics=None, start_time=None, end_time=None):
    """
    Create overview section content for the document.
    Returns tuple of (overview_content, formatting_info) where formatting_info is a dict with:
    - folder_link_info: (link_start_index, link_end_index, folder_url)
    - bold_labels: list of (start_index, end_index) for "Files Processed:", "Images with Errors:", "Metrics:", "Prompt Used:"
    - prompt_text_range: (start_index, end_index) for the prompt text
    - disclaimer_range: (start_index, end_index) for the disclaimer text (for yellow highlight)
    
    Args:
        pages: List of page dictionaries with 'name', 'webViewLink', 'text'
        config: Configuration dictionary
        prompt_text: The prompt text used for transcription
        metrics: Optional dictionary with calculated metrics (total_time, avg_time_per_page, etc.)
        start_time: Optional datetime object for when transcription started
        end_time: Optional datetime object for when transcription ended
    """
    # Handle both normalized (nested) and legacy (flat) config formats
    googlecloud_config = config.get('googlecloud', {})
    drive_folder_id = googlecloud_config.get('drive_folder_id') or config.get('drive_folder_id')
    if not drive_folder_id:
        raise KeyError("'drive_folder_id' not found in config (checked googlecloud.drive_folder_id and top-level)")
    document_name = googlecloud_config.get('document_name') or config.get('document_name', 'Unknown')
    archive_index = config.get('archive_index')
    ocr_model_id = googlecloud_config.get('ocr_model_id') or config.get('ocr_model_id')
    prompt_file = config.get('prompt_file')
    
    # Get folder link from the first page
    folder_link = pages[0]['webViewLink'] if pages else ""
    # Extract folder ID from the link for a cleaner folder link
    if 'folders/' in folder_link:
        folder_id = folder_link.split('folders/')[1].split('/')[0]
        folder_url = f"https://drive.google.com/drive/folders/{folder_id}"
    else:
        # Use drive_folder_id to construct folder URL if webViewLink doesn't have folder info
        folder_url = f"https://drive.google.com/drive/folders/{drive_folder_id}"
    
    # Count successful and failed transcriptions
    successful_pages = [p for p in pages if p['text'] and not p['text'].startswith('[Error')]
    failed_pages = [p for p in pages if not p['text'] or p['text'].startswith('[Error')]
    
    # Get file range
    if pages:
        start_file = pages[0]['name']
        end_file = pages[-1]['name']
        file_count = len(pages)
    else:
        start_file = "N/A"
        end_file = "N/A"
        file_count = 0
    
    # Disclaimer text in 3 languages
    disclaimer_text = """Нейросеть допускает много неточностей в переводе имен и фамилий - использовать как приблизительный перевод рукописного текста и перепроверять с источником!

Нейромережа допускає багато неточностей у перекладі імен та прізвищ - використовувати як приблизний переклад рукописного тексту та перевіряти з джерелом!

The neural network makes many inaccuracies in translating names and surnames - use as an approximate translation of handwritten text and verify with the source!"""
    
    # Create overview content with disclaimer first, then document name as link text
    overview_content = f"""TRANSCRIPTION RUN SUMMARY

{disclaimer_text}

Name: {document_name}
Folder Link: {document_name}
"""
    
    # Calculate folder link position (after "Folder Link: ")
    folder_link_start = overview_content.find("Folder Link: ") + len("Folder Link: ")
    folder_link_end = folder_link_start + len(document_name)
    folder_link_info = (folder_link_start, folder_link_end, folder_url)
    
    # Add archive index if available
    if archive_index:
        overview_content += f"Archive Index: {archive_index}\n"
    
    # Add time start and time end
    if start_time:
        start_time_str = start_time.strftime("%Y-%m-%d %H:%M:%S")
        overview_content += f"Time Start: {start_time_str}\n"
    if end_time:
        end_time_str = end_time.strftime("%Y-%m-%d %H:%M:%S")
        overview_content += f"Time End: {end_time_str}\n"
    
    # Track positions for bold labels
    files_processed_label = "Files Processed:"
    images_errors_label = f"Images with Errors ({len(failed_pages)}):"
    metrics_label = "Metrics:"
    prompt_used_label = "Prompt Used:"
    
    overview_content += f"""Model: {ocr_model_id}
Prompt File: {prompt_file}

{files_processed_label}
Count: {file_count}
Start: {start_file}
End: {end_file}

{images_errors_label}
"""
    
    # Calculate positions for bold labels
    files_processed_start = overview_content.find(files_processed_label)
    files_processed_end = files_processed_start + len(files_processed_label)
    
    images_errors_start = overview_content.find(images_errors_label)
    images_errors_end = images_errors_start + len(images_errors_label)
    
    # Add failed images list
    if failed_pages:
        for page in failed_pages:
            overview_content += f"- {page['name']}\n"
    else:
        overview_content += "None\n"
    
    # Add metrics section
    if metrics:
        overview_content += f"""
{metrics_label}
Total Time: {metrics['total_time']:.1f} seconds
Average Time per Page (secs): {metrics['avg_time_per_page']:.2f}
Total Input Tokens Used: {metrics['total_input_tokens']:,}
Total Output Tokens Used: {metrics['total_output_tokens']:,}
Estimated Cost per Run: ${metrics['estimated_cost_per_run']:.6f}
Estimated Cost Per Page: ${metrics['estimated_cost_per_page']:.6f}
"""
    else:
        overview_content += f"""
{metrics_label}
Total Time: N/A
Average Time per Page (secs): N/A
Total Input Tokens Used: N/A
Total Output Tokens Used: N/A
Estimated Cost per Run: N/A
Estimated Cost Per Page: N/A
"""
    
    # Calculate metrics label position for bold formatting
    metrics_start = overview_content.find(metrics_label)
    metrics_end = metrics_start + len(metrics_label)
    
    # Calculate prompt text position
    prompt_used_start = overview_content.find(prompt_used_label)
    prompt_used_end = prompt_used_start + len(prompt_used_label)
    
    overview_content += f"""
{prompt_used_label}
{prompt_text}

{'='*50}

"""
    
    # Calculate prompt text range (after the label and newline)
    prompt_text_start = overview_content.find(prompt_used_label) + len(prompt_used_label) + 1  # +1 for newline
    prompt_text_end = prompt_text_start + len(prompt_text)
    
    # Recalculate disclaimer range (positions may have shifted)
    disclaimer_start = overview_content.find(disclaimer_text.split('\n')[0])
    disclaimer_end = disclaimer_start + len(disclaimer_text)
    
    formatting_info = {
        'folder_link_info': folder_link_info,
        'bold_labels': [
            (files_processed_start, files_processed_end),
            (images_errors_start, images_errors_end),
            (metrics_start, metrics_end),
            (prompt_used_start, prompt_used_end)
        ],
        'prompt_text_range': (prompt_text_start, prompt_text_end),
        'disclaimer_range': (disclaimer_start, disclaimer_end)
    }
    
    return overview_content, formatting_info


def create_local_overview_section(pages, config: dict, prompt_text: str, metrics=None, start_time=None, end_time=None, error_info=None):
    """
    Create overview section for local mode (simplified version without Google Drive links).
    
    Args:
        pages: List of page dictionaries with 'name', 'webViewLink', 'text'
        config: Configuration dictionary
        prompt_text: The prompt text used for transcription
        metrics: Optional dictionary with calculated metrics
        start_time: Optional datetime for session start
        end_time: Optional datetime for session end
        error_info: Optional dict with error details {'type': str, 'message': str, 'status_code': int, 'next_image_number': int}
        
    Returns:
        String with overview content
    """
    archive_index = config.get('archive_index', 'Unknown')
    ocr_model_id = config.get('local', {}).get('ocr_model_id', 'gemini-3-flash-preview')
    prompt_file = config.get('prompt_file', 'Unknown')
    image_dir = config.get('local', {}).get('image_dir', 'Unknown')
    
    # Count successful and failed transcriptions
    successful_pages = [p for p in pages if p.get('text') and not p['text'].startswith('[Error')]
    failed_pages = [p for p in pages if not p.get('text') or p['text'].startswith('[Error')]
    
    # Get file range
    if pages:
        start_file = pages[0]['name']
        end_file = pages[-1]['name']
        file_count = len(pages)
    else:
        start_file = "N/A"
        end_file = "N/A"
        file_count = 0
    
    # Disclaimer text
    disclaimer_text = """Нейросеть допускает много неточностей в переводе имен и фамилий - использовать как приблизительный перевод рукописного текста и перепроверять с источником!

Нейромережа допускає багато неточностей у перекладі імен та прізвищ - використовувати як приблизний переклад рукописного тексту та перевіряти з джерелом!

The neural network makes many inaccuracies in translating names and surnames - use as an approximate translation of handwritten text and verify with the source!"""
    
    overview_content = f"""TRANSCRIPTION RUN SUMMARY

{disclaimer_text}

Archive Index: {archive_index}
Image Directory: {image_dir}
Model: {ocr_model_id}
Prompt File: {prompt_file}
"""
    
    # Add time info BEFORE Files Processed section
    if start_time:
        overview_content += f"Time Start: {start_time.strftime('%Y-%m-%d %H:%M:%S')}\n"
    if end_time:
        overview_content += f"Time End: {end_time.strftime('%Y-%m-%d %H:%M:%S')}\n"
    
    # Add error information if present
    if error_info:
        overview_content += f"\n**RUN INTERRUPTED BY ERROR**\n"
        overview_content += f"Error Type: {error_info.get('type', 'Unknown')}\n"
        if error_info.get('status_code'):
            overview_content += f"Error Code: {error_info['status_code']}\n"
        overview_content += f"Error Message: {error_info.get('message', 'Unknown error')}\n"
        if error_info.get('next_image_number'):
            overview_content += f"\nTo resume: Update config 'image_start_number' to {error_info['next_image_number']}\n"
        overview_content += f"\n"
    
    overview_content += f"""
Files Processed:
Count: {file_count}
Start: {start_file}
End: {end_file}

Images with Errors ({len(failed_pages)}):
"""
    
    if failed_pages:
        for page in failed_pages:
            overview_content += f"- {page['name']}\n"
    else:
        overview_content += "None\n"
    
    # Add metrics
    if metrics:
        overview_content += f"""
Metrics:
Total Time: {metrics.get('total_time', 0):.1f} seconds ({metrics.get('total_time', 0)/60:.1f} minutes)
Average Time per Page: {metrics.get('avg_time_per_page', 0):.2f} seconds
"""
        if 'total_input_tokens' in metrics:
            overview_content += f"Total Input Tokens: {metrics['total_input_tokens']:,}\n"
            overview_content += f"Total Output Tokens: {metrics['total_output_tokens']:,}\n"
            if metrics.get('total_cached_tokens', 0) > 0:
                overview_content += f"Total Cached Tokens: {metrics['total_cached_tokens']:,}\n"
            if 'estimated_cost_per_run' in metrics:
                overview_content += f"Estimated Cost per Run: ${metrics['estimated_cost_per_run']:.6f}\n"
            if 'estimated_cost_per_page' in metrics:
                overview_content += f"Estimated Cost per Page: ${metrics['estimated_cost_per_page']:.6f}\n"
            if 'page_count' in metrics:
                overview_content += f"Pages Processed: {metrics['page_count']}\n"
    
    overview_content += f"""
Prompt Used:
{prompt_text}

{'='*50}

"""
    
    return overview_content


def update_overview_section(docs_service, doc_id, pages, config: dict, prompt_text: str, metrics=None, start_time=None, end_time=None):
    """
    Update the overview section in an existing document with final metrics.
    
    This function finds the overview section (between the header and first page transcription)
    and replaces it with updated content that includes final metrics from all batches.
    
    Args:
        docs_service: Google Docs API service
        doc_id: Document ID
        pages: List of all page dictionaries (from all batches)
        config: Configuration dictionary
        prompt_text: The prompt text used for transcription
        metrics: Final metrics dictionary with overall stats
        start_time: Start time of the transcription run
        end_time: End time of the transcription run
    """
    archive_index = config.get('archive_index')  # Optional field
    import time
    from googleapiclient.errors import HttpError
    
    max_retries = 3
    retry_delay = 30  # seconds
    # Exponential backoff timeouts: 1 min, 2 min, 5 min
    timeout_seconds_list = [60, 120, 300]
    
    for attempt in range(max_retries):
        attempt_start_time = time.time()
        try:
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Attempt {attempt + 1}/{max_retries} to update overview section (timeout: {timeout_seconds_list[attempt]/60:.1f} min)...")
            
            # Get current document state
            doc = docs_service.documents().get(documentId=doc_id).execute()
            
            # Find the overview section
            # It's between the header (Heading 1) and the first page transcription (Heading 2 with archive index)
            content = doc['body']['content']
            
            # Find the overview section boundaries
            # Overview starts at "TRANSCRIPTION RUN SUMMARY" and ends before the first page transcription
            overview_start = None
            overview_end = None
            first_page_header_start = None
            
            # First, find the start of "TRANSCRIPTION RUN SUMMARY" heading and first page header
            summary_heading = "TRANSCRIPTION RUN SUMMARY"
            for i, element in enumerate(content):
                if 'paragraph' in element:
                    para = element['paragraph']
                    if 'paragraphStyle' in para and para['paragraphStyle'].get('namedStyleType') == 'HEADING_2':
                        # Get text content of this paragraph
                        text_content = ""
                        for elem in para.get('elements', []):
                            if 'textRun' in elem:
                                text_content += elem['textRun'].get('content', '')
                        
                        # Check if this is the "TRANSCRIPTION RUN SUMMARY" heading
                        if summary_heading in text_content and not overview_start:
                            overview_start = element['startIndex']
                            if first_page_header_start:
                                break
                        # Check if this is the first page header (contains archive index but not summary heading)
                        elif archive_index and archive_index in text_content and summary_heading not in text_content and not first_page_header_start:
                            first_page_header_start = element['startIndex']
                            if overview_start:
                                break
            
            # If we didn't find overview_start by heading, try fallback method
            if not overview_start:
                # Fallback: look for "TRANSCRIPTION RUN SUMMARY" text in document
                doc_text = ""
                char_positions = []  # Track character positions
                current_pos = 0
                for element in content:
                    if 'paragraph' in element:
                        for elem in element['paragraph'].get('elements', []):
                            if 'textRun' in elem:
                                text = elem['textRun'].get('content', '')
                                doc_text += text
                                char_positions.append((current_pos, current_pos + len(text), element['startIndex']))
                                current_pos += len(text)
                
                summary_pos = doc_text.find("TRANSCRIPTION RUN SUMMARY")
                if summary_pos >= 0:
                    # Find the element that contains this position
                    for start_char, end_char, element_start in char_positions:
                        if start_char <= summary_pos < end_char:
                            overview_start = element_start + (summary_pos - start_char)
                            break
            # Find the end of overview section (start of first page transcription) if not already found
            if not first_page_header_start:
                for i, element in enumerate(content):
                    if 'paragraph' in element:
                        para = element['paragraph']
                        if 'paragraphStyle' in para and para['paragraphStyle'].get('namedStyleType') == 'HEADING_2':
                            text_content = ""
                            for elem in para.get('elements', []):
                                if 'textRun' in elem:
                                    text_content += elem['textRun'].get('content', '')
                            # This is the first page header if it contains archive index but not the summary heading
                            if archive_index and archive_index in text_content and summary_heading not in text_content:
                                first_page_header_start = element['startIndex']
                                break
            
            # Set overview_end to first page header start
            if first_page_header_start:
                overview_end = first_page_header_start
            
            if not overview_start or not overview_end:
                logging.warning("Could not locate overview section boundaries. Skipping overview update.")
                return False
            
            # Prepare new overview content with final metrics
            overview_content, formatting_info = create_overview_section(pages, config, prompt_text, metrics, start_time, end_time)
            folder_link_info = formatting_info['folder_link_info']
            bold_labels = formatting_info['bold_labels']
            prompt_text_range = formatting_info['prompt_text_range']
            disclaimer_range = formatting_info.get('disclaimer_range', (0, 0))
            folder_link_start_offset, folder_link_end_offset, folder_url = folder_link_info
            
            # Calculate position of "TRANSCRIPTION RUN SUMMARY" heading
            summary_heading = "TRANSCRIPTION RUN SUMMARY"
            summary_heading_start = 0
            summary_heading_end = len(summary_heading)
            
            # Delete old overview and insert new one
            update_requests = [
            {
                'deleteContentRange': {
                    'range': {
                        'startIndex': overview_start,
                        'endIndex': overview_end
                    }
                }
            },
            {
                'insertText': {
                    'location': {'index': overview_start},
                    'text': overview_content
                }
            },
            {
                'updateParagraphStyle': {
                    'range': {'startIndex': overview_start, 'endIndex': overview_start + len(overview_content)},
                    'paragraphStyle': {
                        'namedStyleType': 'NORMAL_TEXT',
                        'alignment': 'START'
                    },
                    'fields': 'namedStyleType,alignment'
                }
            },
            {
                'updateParagraphStyle': {
                    'range': {'startIndex': overview_start + summary_heading_start, 'endIndex': overview_start + summary_heading_end + 1},
                    'paragraphStyle': {
                        'namedStyleType': 'HEADING_2',
                        'alignment': 'START'
                    },
                    'fields': 'namedStyleType,alignment'
                }
            },
            {
                'updateTextStyle': {
                    'range': {'startIndex': overview_start + summary_heading_start, 'endIndex': overview_start + summary_heading_end},
                    'textStyle': {
                        'bold': False
                    },
                    'fields': 'bold'
                }
            }
            ]
            
            # Add disclaimer highlight (yellow background)
            if disclaimer_range[1] > disclaimer_range[0]:
                update_requests.append({
                    'updateTextStyle': {
                        'range': {
                            'startIndex': overview_start + disclaimer_range[0],
                            'endIndex': overview_start + disclaimer_range[1]
                        },
                        'textStyle': {
                            'backgroundColor': {
                                'color': {
                                    'rgbColor': {
                                        'red': 1.0,
                                        'green': 1.0,
                                        'blue': 0.0
                                    }
                                }
                            }
                        },
                        'fields': 'backgroundColor'
                    }
                })
            
            # Add folder link
            update_requests.append({
                'updateTextStyle': {
                    'range': {
                        'startIndex': overview_start + folder_link_start_offset,
                        'endIndex': overview_start + folder_link_end_offset
                    },
                    'textStyle': {
                        'link': {
                            'url': folder_url
                        }
                    },
                    'fields': 'link'
                }
            })
            
            # Add bold formatting for labels
            for label_start, label_end in bold_labels:
                update_requests.append({
                    'updateTextStyle': {
                        'range': {
                            'startIndex': overview_start + label_start,
                            'endIndex': overview_start + label_end
                        },
                        'textStyle': {
                            'bold': True
                        },
                        'fields': 'bold'
                    }
                })
            
            # Format prompt text (6pt Roboto Mono Normal)
            if prompt_text_range[1] > prompt_text_range[0]:
                update_requests.append({
                    'updateTextStyle': {
                        'range': {
                            'startIndex': overview_start + prompt_text_range[0],
                            'endIndex': overview_start + prompt_text_range[1]
                        },
                        'textStyle': {
                            'fontSize': {
                                'magnitude': 6,
                                'unit': 'PT'
                            },
                            'weightedFontFamily': {
                                'fontFamily': 'Roboto Mono'
                            }
                        },
                        'fields': 'fontSize,weightedFontFamily'
                    }
                })
        
            # Execute update
            docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': update_requests}).execute()
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Overview section updated with final metrics successfully.")
            return True  # Success, exit retry loop
            
        except (TimeoutError, HttpError, ConnectionError, OSError) as e:
            attempt_elapsed = time.time() - attempt_start_time
            error_type = type(e).__name__
            error_msg = f"Attempt {attempt + 1}/{max_retries} failed to update overview section after {attempt_elapsed:.1f}s: {error_type}: {str(e)}"
            logging.warning(f"[{datetime.now().strftime('%H:%M:%S')}] {error_msg}")
            
            if attempt < max_retries - 1:
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Retrying in {retry_delay} seconds... (exponential backoff)")
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
            else:
                # All retries exhausted
                logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] All {max_retries} attempts failed to update overview section: {error_type}: {str(e)}")
                logging.error(f"Full traceback:\n{traceback.format_exc()}")
                return False  # Give up after all retries
                
        except Exception as e:
            # Non-retryable error, log and return
            error_type = type(e).__name__
            logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] Error updating overview section: {error_type}: {str(e)}")
            logging.error(f"Full traceback:\n{traceback.format_exc()}")
            return False
    
    # If we get here, all retries failed
    return False


def add_record_links_to_text(text, archive_index, page_number, web_view_link):
    """
    Find lines starting with ### or #### (like "### Запись 1" or "#### Record 1") and append clickable archive reference.
    For example: "### Запись 1" becomes "### Запись 1 ф201оп4спр104стр22"
    or "#### Record 1" becomes "#### Record 1 ф201оп4спр104стр22"
    where the archive part is a hyperlink.
    
    Returns tuple of (modified_text, link_insertions) where link_insertions is a list of
    (start_index, end_index, url) tuples for creating links in Google Docs.
    """
    if not text or not archive_index or not web_view_link:
        return text, []
    
    import re
    
    lines = text.split('\n')
    modified_lines = []
    link_insertions = []
    current_pos = 0
    
    # Pattern to match ### or #### at start of line (with optional whitespace) followed by any text
    # Matches lines like "### Запись 1", "#### Record 1", "### Запис 1", etc.
    pattern = re.compile(r'^(#{3,4}\s+[^\n]+)', re.IGNORECASE)
    
    for line in lines:
        match = pattern.match(line)
        if match:
            # Found a record header line
            original_line = match.group(1)
            archive_ref = f"{archive_index}стр{page_number}"
            modified_line = f"{original_line} {archive_ref}"
            modified_lines.append(modified_line)
            
            # Calculate position for link insertion (the archive_ref part)
            link_start = current_pos + len(original_line) + 1  # +1 for the space
            link_end = link_start + len(archive_ref)
            link_insertions.append((link_start, link_end, web_view_link))
            
            current_pos += len(modified_line) + 1  # +1 for newline
        else:
            modified_lines.append(line)
            current_pos += len(line) + 1  # +1 for newline
    
    modified_text = '\n'.join(modified_lines)
    return modified_text, link_insertions


def save_transcription_locally(pages, doc_name, config: dict, prompt_text: str, logs_dir: str, metrics=None, start_time=None, end_time=None):
    """
    Save transcription to a local text file when Google Doc creation fails.
    """
    # Handle both normalized (nested) and legacy (flat) config formats
    googlecloud_config = config.get('googlecloud', {})
    document_name = googlecloud_config.get('document_name') or config.get('document_name', 'Unknown')
    
    try:
        # Create output directory if it doesn't exist
        output_dir = os.path.join(logs_dir, "local_transcriptions")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Generate filename
        safe_doc_name = "".join(c for c in doc_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_doc_name = safe_doc_name.replace(' ', '_')
        output_file = os.path.join(output_dir, f"{safe_doc_name}.txt")
        
        # Create overview content (for local files, we just need the text, not formatting info)
        overview_content, _ = create_overview_section(pages, config, prompt_text, metrics, start_time, end_time)
        
        # Write to file
        with open(output_file, 'w', encoding='utf-8') as f:
            # Write document header
            run_date = datetime.now().strftime("%Y%m%d")
            document_header = f"{document_name} {run_date}"
            f.write(document_header + "\n\n")
            
            # Write overview
            f.write(overview_content)
            f.write("\n" + "="*80 + "\n\n")
            
            # Write each page
            for page in pages:
                f.write(f"\n{'='*80}\n")
                f.write(f"FILE: {page['name']}\n")
                f.write(f"SOURCE: {page['webViewLink']}\n")
                f.write(f"{'='*80}\n\n")
                f.write(page['text'] if page['text'] else "[No transcription available]")
                f.write("\n\n")
        
        logging.info(f"Transcription saved locally to: {output_file}")
        return output_file
    except Exception as e:
        logging.error(f"Error saving transcription locally: {str(e)}")
        raise


def write_to_doc(docs_service, drive_service, doc_id, pages, config: dict, prompt_text: str, start_idx=0, metrics=None, start_time=None, end_time=None, write_overview=True, genai_client=None):
    """
    Write transcribed content to a Google Doc using Atomic Page Writes.
    Fetches the true document end index before every page write to prevent index drift errors.
    
    Formatting includes:
    - Overview section with metadata (if write_overview=True)
    - Archive reference + page number as Heading 2 (e.g., "ф201оп4спр104стр22")
    - Source image link
    - Raw Vertex AI output with clickable links on record headers
    
    Args:
        docs_service: Google Docs API service
        drive_service: Google Drive API service
        doc_id: Document ID
        pages: List of page dictionaries
        config: Configuration dictionary
        prompt_text: The prompt text used for transcription
        start_idx: Starting index in pages list (for incremental writes)
        metrics: Optional metrics dictionary
        start_time: Optional start time
        end_time: Optional end time
        write_overview: If True, write overview section and document header (default: True)
        genai_client: Optional Vertex AI Gemini client (needed for title page transcription)
    """
    # Handle both normalized (nested) and legacy (flat) config formats
    googlecloud_config = config.get('googlecloud', {})
    document_name = googlecloud_config.get('document_name') or config.get('document_name', 'Unknown')
    archive_index = config.get('archive_index')  # Optional field
    
    logging.info(f"Preparing document content...")
    if archive_index:
        logging.info(f"Using archive index: {archive_index}")
    
    try:
        # Circuit breaker to prevent cascading failures
        MAX_CONSECUTIVE_FAILURES = 5
        consecutive_failures = 0
        
        # --- PHASE 1: Write Overview (if needed) ---
        if write_overview:
            # Fetch current doc state - start at beginning
            doc = docs_service.documents().get(documentId=doc_id).execute()
            idx = 1  # Start of document
            
            # Prepare Header
            run_date = datetime.now().strftime("%Y%m%d")
            document_header = f"{document_name} {run_date}"
            
            # Insert Header
            header_requests = [
                {
                    'insertText': {
                        'location': {'index': idx},
                        'text': document_header + "\n\n"
                    }
                },
                {
                    'updateParagraphStyle': {
                        'range': {'startIndex': idx, 'endIndex': idx + len(document_header) + 2},
                        'paragraphStyle': {
                            'namedStyleType': 'HEADING_1',
                            'alignment': 'START'
                        },
                        'fields': 'namedStyleType,alignment'
                    }
                }
            ]
            docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': header_requests}).execute()
            logging.info(f"Document header added: {document_header}")
            
            # Refresh Index after header insertion
            doc = docs_service.documents().get(documentId=doc_id).execute()
            # Google Docs content list includes a final EOF character, so endIndex - 1 is the actual end
            idx = doc['body']['content'][-1]['endIndex'] - 1
            
            # Insert title page image and transcribe it if specified
            title_page_transcription = None
            if config.get('title_page_filename') and genai_client:
                idx, title_page_transcription = insert_title_page_image_and_transcribe(
                    docs_service, drive_service, doc_id, config, idx, genai_client, prompt_text
                )
                # Refresh index after image insertion
                doc = docs_service.documents().get(documentId=doc_id).execute()
                idx = doc['body']['content'][-1]['endIndex'] - 1
                
                # Insert title page transcription if available, before overview section
                if title_page_transcription:
                    transcription_text = f"\n\n{title_page_transcription}\n\n"
                    transcription_requests = [{
                        'insertText': {
                            'location': {'index': idx},
                            'text': transcription_text
                        }
                    }]
                    docs_service.documents().batchUpdate(
                        documentId=doc_id,
                        body={'requests': transcription_requests}
                    ).execute()
                    logging.info(f"Title page transcription inserted successfully")
                    # Refresh index after transcription insertion
                    doc = docs_service.documents().get(documentId=doc_id).execute()
                    idx = doc['body']['content'][-1]['endIndex'] - 1
            
            # CRITICAL: Refresh index right before inserting overview to ensure accuracy
            # This prevents any issues with image insertion affecting index calculations
            doc = docs_service.documents().get(documentId=doc_id).execute()
            idx = doc['body']['content'][-1]['endIndex'] - 1
            
            # Prepare Overview Content
            overview_content, formatting_info = create_overview_section(pages, config, prompt_text, metrics, start_time, end_time)
            folder_link_info = formatting_info['folder_link_info']
            bold_labels = formatting_info['bold_labels']
            prompt_text_range = formatting_info['prompt_text_range']
            disclaimer_range = formatting_info.get('disclaimer_range', (0, 0))
            folder_link_start_offset, folder_link_end_offset, folder_url = folder_link_info
            
            # Calculate position of "TRANSCRIPTION RUN SUMMARY" heading
            summary_heading = "TRANSCRIPTION RUN SUMMARY"
            summary_heading_start = 0
            summary_heading_end = len(summary_heading)
            
            # Construct Overview Requests
            overview_requests = [
            {
                'insertText': {
                    'location': {'index': idx},
                    'text': overview_content
                }
            },
            {
                'updateParagraphStyle': {
                    # Start range after the inserted text to avoid affecting any preceding content (like images)
                    'range': {'startIndex': idx + 1, 'endIndex': idx + len(overview_content)},
                    'paragraphStyle': {
                        'namedStyleType': 'NORMAL_TEXT',
                        'alignment': 'START'
                    },
                    'fields': 'namedStyleType,alignment'
                }
            },
            {
                'updateParagraphStyle': {
                        'range': {'startIndex': idx + summary_heading_start, 'endIndex': idx + summary_heading_end + 1},  # +1 for newline
                        'paragraphStyle': {
                            'namedStyleType': 'HEADING_2',
                            'alignment': 'START'
                        },
                        'fields': 'namedStyleType,alignment'
                    }
                },
                {
                    'updateTextStyle': {
                        'range': {'startIndex': idx + summary_heading_start, 'endIndex': idx + summary_heading_end},
                        'textStyle': {
                            'bold': False
                        },
                        'fields': 'bold'
                    }
                }
            ]
            
            # Add folder link styling only if range is valid (not empty)
            if folder_link_end_offset > folder_link_start_offset:
                overview_requests.append({
                    'updateTextStyle': {
                        'range': {'startIndex': idx + folder_link_start_offset, 'endIndex': idx + folder_link_end_offset},
                        'textStyle': {
                            'link': {'url': folder_url},
                            'foregroundColor': {'color': {'rgbColor': {'red': 0.0, 'green': 0.0, 'blue': 1.0}}},
                            'underline': True
                        },
                        'fields': 'link,foregroundColor,underline'
                    }
                })
            
            # Add bold formatting for labels (only if ranges are valid)
            for label_start, label_end in bold_labels:
                if label_end > label_start:  # Validate range is not empty
                    overview_requests.append({
                        'updateTextStyle': {
                            'range': {'startIndex': idx + label_start, 'endIndex': idx + label_end},
                            'textStyle': {
                                'bold': True
                            },
                            'fields': 'bold'
                        }
                    })
            
            # Add yellow highlight for disclaimer (only if range is valid)
            if disclaimer_range[1] > disclaimer_range[0]:
                disclaimer_start, disclaimer_end = disclaimer_range
                overview_requests.append({
                    'updateTextStyle': {
                        'range': {'startIndex': idx + disclaimer_start, 'endIndex': idx + disclaimer_end},
                        'textStyle': {
                            'backgroundColor': {
                                'color': {
                                    'rgbColor': {
                                        'red': 1.0,
                                        'green': 1.0,
                                        'blue': 0.0
                                    }
                                }
                            }
                        },
                        'fields': 'backgroundColor'
                    }
                })
            
            # Add formatting for prompt text (6pt Roboto Mono) - only if range is valid
            prompt_text_start, prompt_text_end = prompt_text_range
            if prompt_text_end > prompt_text_start:  # Validate range is not empty
                overview_requests.append({
                    'updateTextStyle': {
                        'range': {'startIndex': idx + prompt_text_start, 'endIndex': idx + prompt_text_end},
                        'textStyle': {
                            'fontSize': {
                                'magnitude': 6.0,
                                'unit': 'PT'
                            },
                            'weightedFontFamily': {
                                'fontFamily': 'Roboto Mono',
                                'weight': 400  # Normal weight
                            }
                        },
                        'fields': 'fontSize,weightedFontFamily'
                    }
                })
            
            # Execute Overview
            docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': overview_requests}).execute()
            consecutive_failures = 0  # Reset counter on success
            logging.info("Overview section added successfully.")
        
        # --- PHASE 2: Write Page Transcriptions (Atomic Writes - One by One) ---
        for i, item in enumerate(pages[start_idx:], start=start_idx + 1):
            try:
                # CRITICAL FIX: Get Fresh Index Before Every Page Write
                # This ensures our index is 100% accurate and prevents "Precondition check failed" errors
                doc = docs_service.documents().get(documentId=doc_id).execute()
                current_idx = doc['body']['content'][-1]['endIndex'] - 1
                
                page_requests = []
                
                # Prepare page data
                page_number = i
                if archive_index:
                    page_header = f"{archive_index}стр{page_number}"
                else:
                    page_header = item['name']
                
                link_text = f"Src Img Url: {item['name']}"
                
                # Build requests for this page atomically
                # 1. Insert Header
                page_requests.append({
                        'insertText': {
                        'location': {'index': current_idx},
                            'text': f"{page_header}\n"
                        }
                })
                page_requests.append({
                        'updateParagraphStyle': {
                        'range': {'startIndex': current_idx, 'endIndex': current_idx + len(page_header) + 1},
                            'paragraphStyle': {
                            'namedStyleType': 'HEADING_2',
                                'alignment': 'START'
                            },
                            'fields': 'namedStyleType,alignment'
                        }
                })
                current_idx += len(page_header) + 1
                
                # 2. Insert Image Link
                page_requests.append({
                        'insertText': {
                        'location': {'index': current_idx},
                            'text': link_text + "\n"
                        }
                })
                link_val_start = current_idx + len("Src Img Url: ")
                link_val_end = current_idx + len(link_text)
                page_requests.append({
                        'updateTextStyle': {
                        'range': {'startIndex': link_val_start, 'endIndex': link_val_end},
                            'textStyle': {
                                'link': {'url': item['webViewLink']},
                                'foregroundColor': {'color': {'rgbColor': {'red': 0.0, 'green': 0.0, 'blue': 1.0}}},
                                'underline': True
                            },
                            'fields': 'link,foregroundColor,underline'
                        }
                })
                current_idx += len(link_text) + 1
                
                # 3. Insert Transcription Body
                if item['text']:
                    modified_text, link_insertions = add_record_links_to_text(
                        item['text'], 
                        archive_index, 
                        page_number, 
                        item['webViewLink']
                    )
                    text_to_insert = modified_text + "\n\n"
                    body_start_idx = current_idx
                    
                    page_requests.append({
                        'insertText': {
                            'location': {'index': current_idx},
                            'text': text_to_insert
                        }
                    })
                    page_requests.append({
                        'updateParagraphStyle': {
                            'range': {'startIndex': current_idx, 'endIndex': current_idx + len(text_to_insert)},
                            'paragraphStyle': {
                                'namedStyleType': 'NORMAL_TEXT',
                                'alignment': 'START'
                            },
                            'fields': 'namedStyleType,alignment'
                        }
                    })
                    
                    # Apply links to ### record headers
                    for l_start, l_end, l_url in link_insertions:
                        abs_start = body_start_idx + l_start
                        abs_end = body_start_idx + l_end
                        page_requests.append({
                            'updateTextStyle': {
                                'range': {'startIndex': abs_start, 'endIndex': abs_end},
                                'textStyle': {
                                    'link': {'url': l_url},
                                    'foregroundColor': {'color': {'rgbColor': {'red': 0.0, 'green': 0.0, 'blue': 1.0}}},
                                    'underline': True
                                },
                                'fields': 'link,foregroundColor,underline'
                            }
                        })
                    
                # 4. Execute immediately for this page (Atomic Write)
                if page_requests:
                    docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': page_requests}).execute()
                    consecutive_failures = 0  # Reset counter on success
                logging.info(f"Added transcription for '{item['name']}' to document (header: {page_header})")
                
            except Exception as e:
                consecutive_failures += 1
                logging.error(f"Error writing page {i} ('{item['name']}') to Doc: {str(e)}")
                logging.warning(f"Consecutive failures: {consecutive_failures}/{MAX_CONSECUTIVE_FAILURES}")
                
                # Circuit breaker: stop if too many consecutive failures
                if consecutive_failures >= MAX_CONSECUTIVE_FAILURES:
                    logging.error(f"Circuit breaker triggered: {consecutive_failures} consecutive failures")
                    logging.error(f"Stopping document write to prevent cascading errors")
                    logging.error(f"Successfully wrote {i - consecutive_failures} out of {len(pages)} pages")
                    logging.error(f"Use recovery_script.py with the AI response log to recover remaining pages")
                    raise Exception(f"Exceeded maximum consecutive failures ({MAX_CONSECUTIVE_FAILURES}). Document write stopped.")
                
                # Try to write an error marker to the doc so the user knows something is missing
                try:
                    doc = docs_service.documents().get(documentId=doc_id).execute()
                    err_idx = doc['body']['content'][-1]['endIndex'] - 1
                    docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': [{
                        'insertText': {
                            'location': {'index': err_idx},
                            'text': f"\n[ERROR WRITING CONTENT FOR {item['name']}: {str(e)}]\n\n"
                        }
                    }]}).execute()
                except Exception as error_write_error:
                    logging.error(f"Could not write error message to document: {error_write_error}")
                    # Continue to next page even if error marker fails
        
        logging.info("Google Doc updated successfully")
        
    except Exception as e:
        logging.error(f"Critical error in write_to_doc: {str(e)}")
        raise


# ------------------------- SHARED PROCESSING LOGIC -------------------------

def process_all_local(images: list, handlers: dict, prompt_text: str, config: dict, ai_logger) -> tuple:
    """
    Process all images in local mode (simpler processing, no batching).
    
    Returns:
        tuple: (transcribed_pages, start_time, end_time, usage_metadata_list, timing_list)
    
    Args:
        images: List of image metadata dictionaries
        handlers: Dictionary of strategy handlers (from ModeFactory)
        prompt_text: Prompt text for transcription
        config: Configuration dictionary
        ai_logger: Logger for AI responses
        
    Returns:
        Tuple of (transcribed_pages, start_time, end_time)
    """
    image_source = handlers['image_source']
    ai_client = handlers['ai_client']
    output = handlers.get('output')  # Get output handler for incremental writing
    
    transcribed_pages = []
    usage_metadata_list = []
    timing_list = []
    start_time = datetime.now()
    last_image_end_time = None
    
    # Cost tracking
    total_prompt_tokens = 0
    total_completion_tokens = 0
    total_tokens = 0
    
    # Gemini 3 Flash pricing (as of January 2026)
    # Pricing for gemini-3-flash-preview model
    # Note: Official Google pricing shows $0.50/$3.00, but this project uses $0.15/$0.60
    # as documented in README.md (may reflect early access pricing, discounts, or project-specific rates)
    # This is different from Gemini 1.5 Flash ($0.075/$0.30)
    PROMPT_COST_PER_1M_TOKENS = 0.15  # $0.15 per 1M input tokens (Gemini 3 Flash)
    COMPLETION_COST_PER_1M_TOKENS = 0.60  # $0.60 per 1M output tokens (Gemini 3 Flash)
    
    total_images = len(images)
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Starting transcription of {total_images} images in LOCAL mode...")
    ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === Starting transcription of {total_images} images (LOCAL mode) ===")
    
    # Set up Rich progress bar
    from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeElapsedColumn, TimeRemainingColumn
    from rich.console import Console
    
    console = Console()
    
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TextColumn("•"),
        TextColumn("[cyan]{task.completed}/{task.total}"),
        TextColumn("•"),
        TimeElapsedColumn(),
        TextColumn("•"),
        TimeRemainingColumn(),
        console=console,
        transient=False
    ) as progress:
        task = progress.add_task(
            f"[cyan]Processing {total_images} images...",
            total=total_images
        )
        
        for global_idx, img_info in enumerate(images, 1):
            image_start_time = datetime.now()
            image_name = img_info['name']
            
            # Log gap detection
            if last_image_end_time:
                gap_seconds = (image_start_time - last_image_end_time).total_seconds()
                if gap_seconds > 60:  # Log if gap is more than 1 minute
                    logging.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: Large time gap detected: {gap_seconds:.1f} seconds ({gap_seconds/60:.1f} minutes) between previous image and '{image_name}'")
                    ai_logger.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: Time gap of {gap_seconds:.1f}s ({gap_seconds/60:.1f} min) before {image_name}")
            
            # Update progress bar
            progress.update(
                task,
                advance=0,  # Don't advance yet, we'll do it after processing
                description=f"[cyan]Processing: {image_name[:50]}...[/cyan]"
            )
            
            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Processing image {global_idx}/{total_images}: '{image_name}'")
            ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === Processing image {global_idx}/{total_images}: {image_name} ===")
            
            try:
                # Get image bytes
                download_start = datetime.now()
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Loading image '{image_name}'...")
                img_bytes = image_source.get_image_bytes(img_info)
                download_elapsed = (datetime.now() - download_start).total_seconds()
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Image '{image_name}' loaded in {download_elapsed:.1f}s, starting transcription...")
                
                # Transcribe image
                transcription_start = datetime.now()
                text, elapsed_time, usage_metadata = ai_client.transcribe(img_bytes, image_name, prompt_text)
                transcription_elapsed = (datetime.now() - transcription_start).total_seconds()
                
                # Check for error responses from transcribe()
                if text is None:
                    text = "[No transcription text received]"
                elif isinstance(text, str) and text.startswith("[Error during transcription:"):
                    # Critical error - stop execution
                    error_msg = text
                    image_end_time = datetime.now()
                    image_total_elapsed = (image_end_time - image_start_time).total_seconds()
                    
                    logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] ✗ Failed to transcribe image {global_idx}/{total_images}: '{image_name}' after {transcription_elapsed:.1f}s")
                    logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] Error: {error_msg}")
                    
                    # Check if it's an API key error - stop immediately
                    if "API key" in error_msg or "API_KEY" in error_msg or "INVALID_ARGUMENT" in error_msg:
                        logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] CRITICAL: Invalid API key detected. Stopping execution.")
                        logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] Please check your API key in the configuration file or GEMINI_API_KEY environment variable.")
                        raise ValueError(f"Invalid API key: {error_msg}")
                    
                    # For other errors, raise exception to stop processing
                    raise RuntimeError(f"Transcription failed for {image_name}: {error_msg}")
                
                # Get image URL for output
                image_url = image_source.get_image_url(img_info)
                
                transcribed_pages.append({
                    'name': image_name,
                    'webViewLink': image_url,
                    'text': text
                })
                
                # Collect metrics
                timing_list.append(elapsed_time)
                usage_metadata_list.append(usage_metadata)
                
                # Track token usage for cost estimation
                current_cost = 0.0
                if usage_metadata:
                    # Handle both dict (LOCAL mode) and Pydantic object (GOOGLECLOUD mode)
                    if isinstance(usage_metadata, dict):
                        prompt_tokens = usage_metadata.get('prompt_tokens', 0) or 0
                        completion_tokens = usage_metadata.get('completion_tokens', 0) or 0
                    else:
                        # GOOGLECLOUD mode: usage_metadata is a Pydantic object
                        prompt_tokens = getattr(usage_metadata, 'prompt_token_count', 0) or 0
                        completion_tokens = getattr(usage_metadata, 'candidates_token_count', 0) or 0
                    
                    total_prompt_tokens += prompt_tokens
                    total_completion_tokens += completion_tokens
                    total_tokens += prompt_tokens + completion_tokens
                    
                    # Calculate current cost
                    current_cost = (total_prompt_tokens / 1_000_000 * PROMPT_COST_PER_1M_TOKENS) + \
                                  (total_completion_tokens / 1_000_000 * COMPLETION_COST_PER_1M_TOKENS)
                
                image_end_time = datetime.now()
                image_total_elapsed = (image_end_time - image_start_time).total_seconds()
                last_image_end_time = image_end_time
                
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Successfully completed image {global_idx}/{total_images}: '{image_name}' (transcription: {transcription_elapsed:.1f}s, total: {image_total_elapsed:.1f}s)")
                ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Completed {image_name} - Transcription: {transcription_elapsed:.1f}s, Total: {image_total_elapsed:.1f}s")
                
                # Update progress bar with cost info
                cost_str = f"Est. cost: ${current_cost:.4f}" if usage_metadata and current_cost > 0 else ""
                progress.update(
                    task,
                    advance=1,
                    description=f"[cyan]Processed: {image_name[:40]}... {cost_str}[/cyan]"
                )
                
                # Write transcription incrementally to log file
                if output:
                    try:
                        # Verify the page has text before writing
                        last_page = transcribed_pages[-1]
                        if not last_page.get('text'):
                            logging.warning(f"Page '{last_page.get('name', 'unknown')}' has no text, skipping write_batch")
                        else:
                            output.write_batch([last_page], batch_num=global_idx, is_first=(global_idx == 1))
                    except Exception as e:
                        logging.error(f"Failed to write transcription incrementally: {e}")
                        import traceback
                        logging.error(traceback.format_exc())
                
                # Log progress
                progress_pct = (global_idx / total_images) * 100
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Progress: {global_idx}/{total_images} images ({progress_pct:.1f}%)")
                
            except (ValueError, RuntimeError) as e:
                # Critical errors (API key, etc.) - stop execution immediately
                image_end_time = datetime.now()
                image_total_elapsed = (image_end_time - image_start_time).total_seconds()
                error_type = type(e).__name__
                
                # Check if this is a 503 Service Unavailable error
                error_str = str(e)
                is_503_error = 'status 503' in error_str or '503' in error_str
                
                if is_503_error and output:
                    # Write user-friendly 503 error message to transcription log
                    current_img_num = extract_image_number(image_name) or global_idx
                    friendly_message = ("[SERVICE TEMPORARILY UNAVAILABLE]\n\n"
                        "The Gemini API service is currently unavailable (HTTP 503). "
                        "This is a temporary issue on Google's side, not a problem with your configuration.\n\n"
                        "What to do next:\n"
                        "1. Wait 5-10 minutes for the service to recover\n"
                        "2. Check Google Cloud Status page: https://status.cloud.google.com/\n"
                        "3. Retry the transcription by running the script again\n"
                        "4. If the issue persists, try again later\n\n"
                        f"To resume from this image when the service is available:\n"
                        f"- Update your config: image_start_number = {current_img_num}\n"
                        "- Then run the script again\n\n"
                        f"Error details: {error_str}")
                    try:
                        image_url = image_source.get_image_url(img_info)
                        output.write_batch([{
                            'name': image_name,
                            'webViewLink': image_url,
                            'text': friendly_message
                        }], batch_num=global_idx, is_first=(global_idx == 1))
                        logging.info(f"Wrote 503 error message to transcription log for {image_name}")
                    except Exception as write_error:
                        logging.warning(f"Failed to write 503 error message to log: {write_error}")
                
                error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] CRITICAL ERROR transcribing image {global_idx}/{total_images} '{image_name}' after {image_total_elapsed:.1f}s: {error_type}: {str(e)}"
                logging.error(error_msg)
                logging.error(f"Full traceback:\n{traceback.format_exc()}")
                ai_logger.error(f"[{datetime.now().strftime('%H:%M:%S')}] CRITICAL ERROR processing {image_name}: {error_type}: {str(e)}")
                ai_logger.error(f"Traceback:\n{traceback.format_exc()}")
                
                # Advance progress bar before re-raising
                progress.update(task, advance=1)
                
                # Re-raise to stop execution
                raise
            except Exception as e:
                image_end_time = datetime.now()
                image_total_elapsed = (image_end_time - image_start_time).total_seconds()
                last_image_end_time = image_end_time
                error_type = type(e).__name__
                
                # Calculate next image number to start from in case of failure
                current_image_number = extract_image_number(image_name)
                if current_image_number is not None:
                    next_image_number = current_image_number + 1
                else:
                    # Fallback: use position-based calculation
                    image_start_number = config.get('image_start_number', 1)
                    next_image_number = image_start_number + global_idx
                
                error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Error transcribing image {global_idx}/{total_images} '{image_name}' after {image_total_elapsed:.1f}s: {error_type}: {str(e)}"
                logging.error(error_msg)
                logging.error(f"Full traceback:\n{traceback.format_exc()}")
                logging.error(f"RESUME INFO: To resume from this point, update config image_start_number = {next_image_number} (filename number from '{image_name}')")
                ai_logger.error(f"[{datetime.now().strftime('%H:%M:%S')}] ERROR processing {image_name}: {error_type}: {str(e)}")
                ai_logger.error(f"RESUME INFO: Update config image_start_number = {next_image_number} to resume from next image (current image filename number: {current_image_number})")
                ai_logger.error(f"Traceback:\n{traceback.format_exc()}")
                
                # Check if this is a 503 Service Unavailable error
                error_str = str(e)
                is_503_error = 'status 503' in error_str or '503' in error_str
                
                # Add error message as text - use friendly message for 503 errors
                image_url = image_source.get_image_url(img_info)
                if is_503_error:
                    friendly_message = ("[SERVICE TEMPORARILY UNAVAILABLE]\n\n"
                        "The Gemini API service is currently unavailable (HTTP 503). "
                        "This is a temporary issue on Google's side, not a problem with your configuration.\n\n"
                        "What to do next:\n"
                        "1. Wait 5-10 minutes for the service to recover\n"
                        "2. Check Google Cloud Status page: https://status.cloud.google.com/\n"
                        "3. Retry the transcription by running the script again\n"
                        "4. If the issue persists, try again later\n\n"
                        f"To resume from this image when the service is available:\n"
                        f"- Update your config: image_start_number = {next_image_number}\n"
                        "- Then run the script again\n\n"
                        f"Error details: {error_str}")
                    error_text = friendly_message
                else:
                    error_text = f"[Error during transcription: {str(e)}]"
                
                transcribed_pages.append({
                    'name': image_name,
                    'webViewLink': image_url,
                    'text': error_text
                })
                
                # Write to log file immediately if 503 error
                if is_503_error and output:
                    try:
                        output.write_batch([transcribed_pages[-1]], batch_num=global_idx, is_first=(global_idx == 1))
                        logging.info(f"Wrote 503 error message to transcription log for {image_name}")
                    except Exception as write_error:
                        logging.warning(f"Failed to write 503 error message to log: {write_error}")
                
                timing_list.append(None)
                usage_metadata_list.append(None)
                
                # Advance progress bar even on error
                progress.update(task, advance=1)
    
    # Record end time
    end_time = datetime.now()
    total_elapsed = (end_time - start_time).total_seconds()
    
    # Calculate final cost
    final_cost = (total_prompt_tokens / 1_000_000 * PROMPT_COST_PER_1M_TOKENS) + \
                 (total_completion_tokens / 1_000_000 * COMPLETION_COST_PER_1M_TOKENS)
    
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Completed all images: {len(transcribed_pages)} images processed in {total_elapsed:.1f} seconds ({total_elapsed/60:.1f} minutes)")
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Token usage: {total_tokens:,} total ({total_prompt_tokens:,} prompt + {total_completion_tokens:,} completion)")
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Estimated cost: ${final_cost:.4f}")
    
    ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === All images transcription completed ===")
    ai_logger.info(f"Total images: {len(transcribed_pages)}")
    ai_logger.info(f"Total time: {total_elapsed:.1f}s ({total_elapsed/60:.1f} min)")
    ai_logger.info(f"Total tokens: {total_tokens:,} ({total_prompt_tokens:,} prompt + {total_completion_tokens:,} completion)")
    ai_logger.info(f"Estimated cost: ${final_cost:.4f}")
    ai_logger.info(f"Start time: {start_time.strftime('%Y-%m-%d %H:%M:%S')}")
    ai_logger.info(f"End time: {end_time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    return transcribed_pages, start_time, end_time, usage_metadata_list, timing_list


def process_batches_googlecloud(images: list, handlers: dict, prompt_text: str, config: dict, ai_logger) -> list:
    """
    Process images in batches for Google Cloud mode (existing batch processing logic).
    
    Args:
        images: List of image metadata dictionaries
        handlers: Dictionary of strategy handlers (from ModeFactory)
        prompt_text: Prompt text for transcription
        config: Configuration dictionary
        ai_logger: Logger for AI responses
        
    Returns:
        List of transcribed pages with metadata
    """
    image_source = handlers['image_source']
    ai_client = handlers['ai_client']
    output = handlers['output']
    docs_service = handlers['docs_service']
    drive_service = handlers['drive_service']
    genai_client = handlers.get('genai_client')
    
    batch_size_for_doc = config.get('batch_size_for_doc', 10)
    # Handle both normalized (nested) and legacy (flat) config formats
    googlecloud_config = config.get('googlecloud', {})
    document_name = googlecloud_config.get('document_name') or config.get('document_name', 'Unknown')
    image_start_number = config.get('image_start_number', 1)
    
    # Process images in batches for incremental document writing
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Starting transcription of {len(images)} images in batches of {batch_size_for_doc}...")
    ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === Starting batch transcription of {len(images)} images (batch size: {batch_size_for_doc}) ===")
    start_time = datetime.now()
    transcribed_pages = []
    usage_metadata_list = []
    timing_list = []
    last_image_end_time = None
    first_batch = True
    
    # Process images in batches
    total_images = len(images)
    num_batches = (total_images + batch_size_for_doc - 1) // batch_size_for_doc  # Ceiling division
    
    # Cost tracking
    total_prompt_tokens = 0
    total_completion_tokens = 0
    total_tokens = 0
    
    # Gemini 3 Flash pricing (as of January 2026)
    # Pricing for gemini-3-flash-preview model
    # Note: Official Google pricing shows $0.50/$3.00, but this project uses $0.15/$0.60
    # as documented in README.md (may reflect early access pricing, discounts, or project-specific rates)
    # This is different from Gemini 1.5 Flash ($0.075/$0.30)
    PROMPT_COST_PER_1M_TOKENS = 0.15  # $0.15 per 1M input tokens (Gemini 3 Flash)
    COMPLETION_COST_PER_1M_TOKENS = 0.60  # $0.60 per 1M output tokens (Gemini 3 Flash)
    
    # Set up Rich progress bar
    from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeElapsedColumn, TimeRemainingColumn
    from rich.console import Console
    
    console = Console()
    
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TextColumn("•"),
        TextColumn("[cyan]{task.completed}/{task.total}"),
        TextColumn("•"),
        TimeElapsedColumn(),
        TextColumn("•"),
        TimeRemainingColumn(),
        console=console,
        transient=False
    ) as progress:
        task = progress.add_task(
            f"[cyan]Processing {total_images} images in {num_batches} batches...",
            total=total_images
        )
        
        try:
            for batch_num in range(num_batches):
                batch_start_idx = batch_num * batch_size_for_doc
                batch_end_idx = min(batch_start_idx + batch_size_for_doc, total_images)
                batch_images = images[batch_start_idx:batch_end_idx]
                batch_size = len(batch_images)
                
                logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] === Processing batch {batch_num + 1}/{num_batches} (images {batch_start_idx + 1}-{batch_end_idx} of {total_images}) ===")
                ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === Batch {batch_num + 1}/{num_batches}: Processing images {batch_start_idx + 1}-{batch_end_idx} ===")
                
                # Track batch-level transcribed pages and metrics
                batch_transcribed_pages = []
                batch_usage_metadata_list = []
                batch_timing_list = []
                
                for batch_idx, img in enumerate(batch_images, 1):
                    global_idx = batch_start_idx + batch_idx
                    image_start_time = datetime.now()
                    image_name = img['name']
                    
                    # Log gap detection
                    if last_image_end_time:
                        gap_seconds = (image_start_time - last_image_end_time).total_seconds()
                        if gap_seconds > 60:  # Log if gap is more than 1 minute
                            logging.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: Large time gap detected: {gap_seconds:.1f} seconds ({gap_seconds/60:.1f} minutes) between previous image and '{image_name}'")
                            ai_logger.warning(f"[{datetime.now().strftime('%H:%M:%S')}] WARNING: Time gap of {gap_seconds:.1f}s ({gap_seconds/60:.1f} min) before {image_name}")
                    
                    # Update progress bar (without advancing)
                    progress.update(
                        task,
                        advance=0,
                        description=f"[cyan]Batch {batch_num + 1}/{num_batches}: {image_name[:40]}...[/cyan]"
                    )
                    
                    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Processing image {global_idx}/{total_images} (batch {batch_num + 1}, item {batch_idx}/{batch_size}): '{image_name}'")
                    ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === Processing image {global_idx}/{total_images}: {image_name} ===")
                    
                    try:
                        download_start = datetime.now()
                        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Downloading image '{image_name}'...")
                        img_bytes = image_source.get_image_bytes(img)
                        download_elapsed = (datetime.now() - download_start).total_seconds()
                        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Image '{image_name}' downloaded in {download_elapsed:.1f}s, starting transcription...")
                        
                        transcription_start = datetime.now()
                        text, elapsed_time, usage_metadata = ai_client.transcribe(img_bytes, image_name, prompt_text)
                        transcription_elapsed = (datetime.now() - transcription_start).total_seconds()
                        
                        # Ensure text is not None
                        if text is None:
                            text = "[No transcription text received]"
                        
                        batch_transcribed_pages.append({
                            'name': img['name'],
                            'webViewLink': img['webViewLink'],
                            'text': text
                        })
                        
                        # Collect metrics
                        batch_timing_list.append(elapsed_time)
                        batch_usage_metadata_list.append(usage_metadata)
                        
                        # Track token usage for cost estimation
                        current_cost = 0.0
                        if usage_metadata:
                            # Handle both dict (LOCAL mode) and Pydantic object (GOOGLECLOUD mode)
                            if isinstance(usage_metadata, dict):
                                prompt_tokens = usage_metadata.get('prompt_tokens', 0) or 0
                                completion_tokens = usage_metadata.get('completion_tokens', 0) or 0
                            else:
                                # GOOGLECLOUD mode: usage_metadata is a Pydantic object (GenerateContentResponseUsageMetadata)
                                # Access attributes directly, not as dict
                                prompt_tokens = getattr(usage_metadata, 'prompt_token_count', 0) or 0
                                completion_tokens = getattr(usage_metadata, 'candidates_token_count', 0) or 0
                            
                            total_prompt_tokens += prompt_tokens
                            total_completion_tokens += completion_tokens
                            total_tokens += prompt_tokens + completion_tokens
                            
                            # Calculate current cost
                            current_cost = (total_prompt_tokens / 1_000_000 * PROMPT_COST_PER_1M_TOKENS) + \
                                          (total_completion_tokens / 1_000_000 * COMPLETION_COST_PER_1M_TOKENS)
                        
                        image_end_time = datetime.now()
                        image_total_elapsed = (image_end_time - image_start_time).total_seconds()
                        last_image_end_time = image_end_time
                        
                        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Successfully completed image {global_idx}/{total_images}: '{image_name}' (transcription: {transcription_elapsed:.1f}s, total: {image_total_elapsed:.1f}s)")
                        ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Completed {image_name} - Transcription: {transcription_elapsed:.1f}s, Total: {image_total_elapsed:.1f}s")
                        
                        # Update progress bar with cost info (advance once per image)
                        cost_str = f"Est. cost: ${current_cost:.4f}" if usage_metadata and current_cost > 0 else ""
                        progress.update(
                            task,
                            advance=1,
                            description=f"[cyan]Batch {batch_num + 1}/{num_batches}: {image_name[:30]}... {cost_str}[/cyan]"
                        )
                        
                        # Log progress
                        progress_pct = (global_idx / total_images) * 100
                        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Progress: {global_idx}/{total_images} images ({progress_pct:.1f}%)")
                        
                    except Exception as e:
                        image_end_time = datetime.now()
                        image_total_elapsed = (image_end_time - image_start_time).total_seconds()
                        last_image_end_time = image_end_time
                        error_type = type(e).__name__
                        
                        # Calculate next image number to start from in case of failure
                        current_image_number = extract_image_number(image_name)
                        if current_image_number is not None:
                            next_image_number = current_image_number + 1
                        else:
                            # Fallback: if we can't extract number, use position-based calculation
                            next_image_number = image_start_number + global_idx
                        
                        error_msg = f"[{datetime.now().strftime('%H:%M:%S')}] Error transcribing image {global_idx}/{total_images} '{image_name}' after {image_total_elapsed:.1f}s: {error_type}: {str(e)}"
                        logging.error(error_msg)
                        logging.error(f"Full traceback:\n{traceback.format_exc()}")
                        logging.error(f"RESUME INFO: To resume from this point, update config image_start_number = {next_image_number} (filename number from '{image_name}')")
                        ai_logger.error(f"[{datetime.now().strftime('%H:%M:%S')}] ERROR processing {image_name}: {error_type}: {str(e)}")
                        ai_logger.error(f"RESUME INFO: Update config image_start_number = {next_image_number} to resume from next image (current image filename number: {current_image_number})")
                        ai_logger.error(f"Traceback:\n{traceback.format_exc()}")
                        
                        # Add error message as text
                        batch_transcribed_pages.append({
                            'name': img['name'],
                            'webViewLink': img['webViewLink'],
                            'text': f"[Error during transcription: {str(e)}]"
                        })
                        # Add None for metrics on error
                        batch_timing_list.append(None)
                        batch_usage_metadata_list.append(None)
                        
                        # Advance progress bar on error (only once, not in success path)
                        progress.update(task, advance=1)
                
                # After batch is transcribed, write to document
                if batch_transcribed_pages:
                    # Accumulate all transcribed pages and metrics
                    transcribed_pages.extend(batch_transcribed_pages)
                    usage_metadata_list.extend(batch_usage_metadata_list)
                    timing_list.extend(batch_timing_list)
                    
                    if first_batch:
                        # Document should already be initialized in main() before processing starts
                        # Just write the first batch with overview (pass all pages so far)
                        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] First batch completed ({len(batch_transcribed_pages)} images). Writing to document with overview...")
                        ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] First batch completed, writing to document...")
                        
                        # Verify document is initialized
                        if not output.doc_id:
                            # Fallback: initialize if somehow not initialized (shouldn't happen)
                            logging.warning("Document not initialized, initializing now...")
                            output.initialize(config)
                        
                        # Write first batch with overview (pass all pages so far)
                        # Update progress bar to show we're writing (but don't advance)
                        progress.update(
                            task,
                            advance=0,
                            description=f"[cyan]Writing batch {batch_num + 1}/{num_batches} to document...[/cyan]"
                        )
                        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Writing first batch ({len(batch_transcribed_pages)} images) to document with overview...")
                        output.write_batch(transcribed_pages, 1, True)
                        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ First batch written to document")
                        first_batch = False
                    else:
                        # Append subsequent batches to existing document
                        if output.doc_id:
                            # Update progress bar to show we're writing (but don't advance)
                            progress.update(
                                task,
                                advance=0,
                                description=f"[cyan]Writing batch {batch_num + 1}/{num_batches} to document...[/cyan]"
                            )
                            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Writing batch {batch_num + 1} ({len(batch_transcribed_pages)} images) to document...")
                            # Pass all transcribed pages so far (write_batch will calculate start_idx)
                            output.write_batch(transcribed_pages, batch_num + 1, False)
                            logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] ✓ Batch {batch_num + 1} written to document")
                        else:
                            # Document creation failed earlier, save locally
                            logging.warning(f"Cannot write batch {batch_num + 1} to document (doc creation failed). Saving locally...")
                            # Append to local file if it exists, or create new one
                            run_date = datetime.now().strftime("%Y%m%d")
                            doc_name = f"{document_name} {run_date}"
                            save_transcription_locally(batch_transcribed_pages, doc_name, config, prompt_text, "logs", None, None, None)
        
        except Exception as batch_error:
            # Log error and resume information
            error_type = type(batch_error).__name__
            images_processed = len(transcribed_pages) if 'transcribed_pages' in locals() else 0
            
            # Calculate next image number from the last successfully processed image
            next_image_number = None
            if images_processed > 0 and 'transcribed_pages' in locals():
                last_image_name = transcribed_pages[-1]['name']
                last_image_number = extract_image_number(last_image_name)
                if last_image_number is not None:
                    next_image_number = last_image_number + 1
                else:
                    # Fallback: use position-based calculation
                    next_image_number = image_start_number + images_processed
            elif images_processed > 0:
                # Fallback if we can't get the last image name
                next_image_number = image_start_number + images_processed
            
            logging.error(f"[{datetime.now().strftime('%H:%M:%S')}] Error processing batch: {error_type}: {str(batch_error)}")
            logging.error(f"RESUME INFO: Processed {images_processed} images successfully before error")
            if next_image_number is not None:
                last_image_info = f" (last processed: {transcribed_pages[-1]['name'] if images_processed > 0 and 'transcribed_pages' in locals() else 'unknown'})"
                logging.error(f"RESUME INFO: To resume from this point, update config image_start_number = {next_image_number}{last_image_info}")
            logging.error(f"Full traceback:\n{traceback.format_exc()}")
            
            ai_logger.error(f"[{datetime.now().strftime('%H:%M:%S')}] === Batch Processing Error ===")
            ai_logger.error(f"Error type: {error_type}")
            ai_logger.error(f"Error message: {str(batch_error)}")
            ai_logger.error(f"Images processed before error: {images_processed}")
            if next_image_number is not None:
                ai_logger.error(f"RESUME INFO: Update config image_start_number = {next_image_number} to resume from next image")
            ai_logger.error(f"Full traceback:\n{traceback.format_exc()}")
            ai_logger.error(f"=== End Batch Processing Error ===")
            
            # Re-raise to be caught by outer exception handler
            raise
    
    # Record end time
    end_time = datetime.now()
    batch_total_elapsed = (end_time - start_time).total_seconds()
    
    # Calculate final cost
    final_cost = (total_prompt_tokens / 1_000_000 * PROMPT_COST_PER_1M_TOKENS) + \
                 (total_completion_tokens / 1_000_000 * COMPLETION_COST_PER_1M_TOKENS)
    
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Completed all batches: {len(transcribed_pages)} images processed in {batch_total_elapsed:.1f} seconds ({batch_total_elapsed/60:.1f} minutes)")
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Token usage: {total_tokens:,} total ({total_prompt_tokens:,} prompt + {total_completion_tokens:,} completion)")
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Estimated cost: ${final_cost:.4f}")
    
    ai_logger.info(f"[{datetime.now().strftime('%H:%M:%S')}] === All batches transcription completed ===")
    ai_logger.info(f"Total images: {len(transcribed_pages)}")
    ai_logger.info(f"Total time: {batch_total_elapsed:.1f}s ({batch_total_elapsed/60:.1f} min)")
    ai_logger.info(f"Total tokens: {total_tokens:,} ({total_prompt_tokens:,} prompt + {total_completion_tokens:,} completion)")
    ai_logger.info(f"Estimated cost: ${final_cost:.4f}")
    ai_logger.info(f"Start time: {start_time.strftime('%Y-%m-%d %H:%M:%S')}")
    ai_logger.info(f"End time: {end_time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Calculate final metrics and finalize output
    # Calculate metrics if we have timing data (even if usage_metadata is missing, we still have timing)
    # Only skip if both lists are empty (no images processed)
    logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Calculating metrics: timing_list has {len(timing_list)} items, usage_metadata_list has {len(usage_metadata_list)} items")
    if timing_list or usage_metadata_list:
        final_metrics = calculate_metrics(usage_metadata_list, timing_list)
        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Metrics calculated: {final_metrics}")
    else:
        logging.warning(f"[{datetime.now().strftime('%H:%M:%S')}] No timing or usage metadata available, skipping metrics calculation")
        final_metrics = None
    
    # Finalize output (update overview section)
    if output.doc_id and len(transcribed_pages) > 0:
        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Finalizing output with overview update...")
        output.finalize(transcribed_pages, final_metrics)
        logging.info(f"[{datetime.now().strftime('%H:%M:%S')}] Output finalized successfully.")
        logging.info(f"Final metrics: {final_metrics}")
    
    # Calculate end time for summary
    end_time = datetime.now()
    start_time = output.start_time if hasattr(output, 'start_time') and output.start_time else None
    
    # Log comprehensive Run Summary
    _log_run_summary(output, transcribed_pages, final_metrics, start_time, end_time, None, 'googlecloud')
    
    return transcribed_pages


def main(config: dict, prompt_text: str, ai_logger, logs_dir: str, log_filename: str = None, ai_log_filename: str = None):
    """
    Main function to process images and create transcription document.
    Now supports dual-mode operation (LOCAL and GOOGLECLOUD) using strategy pattern.
    """
    try:
        # Detect mode
        mode = detect_mode(config)
        logging.info(f"Detected mode: {mode.upper()}")
        
        # Normalize config to internal format
        normalized_config = normalize_config(config, mode)
        
        # Validate config
        is_valid, errors = validate_config(normalized_config, mode)
        if not is_valid:
            error_msg = f"Configuration errors: {', '.join(errors)}"
            logging.error(error_msg)
            raise ValueError(error_msg)
        
        # Create handlers using ModeFactory
        handlers = ModeFactory.create_handlers(mode, normalized_config)
        logging.info(f"Created handlers for {mode.upper()} mode")
        
        # Extract config values for logging (mode-agnostic)
        archive_index = normalized_config.get('archive_index')
        retry_mode = normalized_config.get('retry_mode', False)
        retry_image_list = normalized_config.get('retry_image_list', [])
        max_images = normalized_config.get('max_images')
        image_start_number = normalized_config.get('image_start_number', 1)
        image_count = normalized_config.get('image_count')
        
        # Mode-specific config extraction
        if mode == 'googlecloud':
            project_id = normalized_config.get('googlecloud', {}).get('project_id', 'Unknown')
            ocr_model_id = normalized_config.get('googlecloud', {}).get('ocr_model_id', 'gemini-3-flash-preview')
            batch_size_for_doc = normalized_config.get('batch_size_for_doc', 10)
            
            # Fetch folder name and set document_name if not provided (Google Cloud only)
            drive_service = handlers['drive_service']
            if 'document_name' not in normalized_config.get('googlecloud', {}) or not normalized_config.get('googlecloud', {}).get('document_name'):
                drive_folder_id = normalized_config.get('googlecloud', {}).get('drive_folder_id')
                if drive_folder_id:
                    fetched_folder_name = get_folder_name(drive_service, drive_folder_id)
                    if fetched_folder_name:
                        normalized_config['googlecloud']['document_name'] = fetched_folder_name
                        logging.info(f"Fetched folder name from Drive API: '{fetched_folder_name}'")
                    else:
                        # Fallback to folder ID if fetch fails
                        normalized_config['googlecloud']['document_name'] = f"Folder_{drive_folder_id[:8]}"
                        logging.warning(f"Could not fetch folder name, using fallback: '{normalized_config['googlecloud']['document_name']}'")
            
            document_name = normalized_config.get('googlecloud', {}).get('document_name', 'Unknown')
        else:  # local mode
            project_id = 'LOCAL'
            ocr_model_id = normalized_config.get('local', {}).get('ocr_model_id', 'gemini-3-flash-preview')
            batch_size_for_doc = None  # Not used in local mode
            document_name = normalized_config.get('local', {}).get('image_dir', 'Unknown')
        
        # Log session start
        ai_logger.info(f"=== Transcription Session Started ===")
        ai_logger.info(f"Session timestamp: {datetime.now().isoformat()}")
        ai_logger.info(f"Mode: {mode.upper()}")
        ai_logger.info(f"Project ID: {project_id}")
        ai_logger.info(f"Document: {document_name}")
        ai_logger.info(f"Archive Index: {archive_index if archive_index else 'None'}")
        ai_logger.info(f"Model: {ocr_model_id}")
        ai_logger.info(f"Retry mode: {retry_mode}")
        if retry_mode:
            ai_logger.info(f"Retry images count: {len(retry_image_list)}")
            ai_logger.info(f"Retry images: {retry_image_list}")
        else:
            ai_logger.info(f"Max images: {max_images}")
            ai_logger.info(f"Image start number: {image_start_number}")
            ai_logger.info(f"Image count: {image_count}")
        if batch_size_for_doc:
            ai_logger.info(f"Batch size for doc: {batch_size_for_doc}")
        ai_logger.info(f"=== Session Configuration ===\n")

        # List images using image source strategy
        image_source = handlers['image_source']
        images = image_source.list_images(normalized_config)
        
        if not images:
            if retry_mode:
                logging.error(f"No retry images found from the retry_image_list")
                ai_logger.error(f"No retry images found from list of {len(retry_image_list)} images")
            else:
                logging.error(f"No images found for the specified range (start: {image_start_number}, count: {image_count})")
                ai_logger.error(f"No images found for range {image_start_number} to {image_start_number + image_count - 1}")
            return
        
        # Initialize output
        output = handlers['output']
        # Pass prompt_text to initialize for LOCAL mode to include it in the log
        if mode == 'local':
            output_id = output.initialize(normalized_config, prompt_text)
        else:
            output_id = output.initialize(normalized_config)
        logging.info(f"Output initialized: {output_id}")
        
        # Variables for finalization
        transcribed_pages = []
        start_time = None
        end_time = None
        error_info = None
        caught_exception = None
        usage_metadata_list = []
        timing_list = []
        
        try:
            # Process images using mode-specific processing function
            if mode == 'googlecloud':
                transcribed_pages = process_batches_googlecloud(images, handlers, prompt_text, normalized_config, ai_logger)
                # For GOOGLECLOUD mode, metrics are calculated inside process_batches_googlecloud
                
                # Upload log files to Google Drive if save_logs_to_source is enabled
                save_logs_to_source = normalized_config.get('googlecloud', {}).get('save_logs_to_source', False)
                if save_logs_to_source:
                    drive_service = handlers.get('drive_service')
                    drive_folder_id = normalized_config.get('googlecloud', {}).get('drive_folder_id')
                    if drive_service and drive_folder_id:
                        try:
                            # Upload main log file
                            if log_filename and os.path.exists(log_filename):
                                upload_log_file_to_drive(drive_service, log_filename, drive_folder_id)
                            # Upload AI responses log file
                            if ai_log_filename and os.path.exists(ai_log_filename):
                                upload_log_file_to_drive(drive_service, ai_log_filename, drive_folder_id)
                        except Exception as upload_error:
                            logging.warning(f"Failed to upload log files to Google Drive: {upload_error}")
            else:  # local mode
                transcribed_pages, start_time, end_time, usage_metadata_list, timing_list = process_all_local(images, handlers, prompt_text, normalized_config, ai_logger)
            
            # Log session completion (to ai_logger)
            ai_logger.info(f"=== Transcription Session Completed ===")
            ai_logger.info(f"Session end timestamp: {datetime.now().isoformat()}")
            ai_logger.info(f"Total images processed: {len(transcribed_pages)}")
            ai_logger.info(f"Successful transcriptions: {len([p for p in transcribed_pages if p['text'] and not p['text'].startswith('[Error')])}")
            ai_logger.info(f"Failed transcriptions: {len([p for p in transcribed_pages if not p['text'] or p['text'].startswith('[Error')])}")
            ai_logger.info(f"=== Session Summary ===\n")
            
            # For GOOGLECLOUD mode, Run Summary is already logged in process_batches_googlecloud
            # For LOCAL mode, Run Summary will be logged in the finally block after finalization
            
        except Exception as e:
            # Save the exception for re-raising after finalization
            caught_exception = e
            
            # Capture error information for finalization
            error_type = type(e).__name__
            error_message = str(e)
            
            # Extract status code if it's an API error
            status_code = None
            if 'status' in error_message.lower():
                import re
                status_match = re.search(r'status (\d+)', error_message, re.IGNORECASE)
                if status_match:
                    status_code = int(status_match.group(1))
            
            # Calculate next image number for resume
            next_image_number = None
            if transcribed_pages:
                last_image_name = transcribed_pages[-1]['name']
                last_image_number = extract_image_number(last_image_name)
                if last_image_number is not None:
                    next_image_number = last_image_number + 1
                else:
                    next_image_number = image_start_number + len(transcribed_pages)
            
            error_info = {
                'type': error_type,
                'message': error_message,
                'status_code': status_code,
                'next_image_number': next_image_number
            }
            
            # Log session error with resume information
            ai_logger.error(f"=== Transcription Session Error ===")
            ai_logger.error(f"Error timestamp: {datetime.now().isoformat()}")
            ai_logger.error(f"Error type: {error_type}")
            ai_logger.error(f"Error: {error_message}")
            ai_logger.error(f"Images processed before error: {len(transcribed_pages)}")
            if next_image_number is not None:
                last_image_info = f" (last processed: {transcribed_pages[-1]['name'] if transcribed_pages else 'unknown'})"
                ai_logger.error(f"RESUME INFO: Update config image_start_number = {next_image_number} to resume from next image{last_image_info}")
            ai_logger.error(f"=== Session Error End ===\n")
            
            logging.error(f"Error in main: {error_type}: {error_message}")
            if next_image_number is not None:
                last_image_info = f" (last processed: {transcribed_pages[-1]['name'] if transcribed_pages else 'unknown'})"
                logging.error(f"RESUME INFO: Processed {len(transcribed_pages)} images successfully before error")
                logging.error(f"RESUME INFO: To resume from this point, update config image_start_number = {next_image_number}{last_image_info}")
            logging.error(f"Full traceback:\n{traceback.format_exc()}")
            
            # Don't re-raise yet, we need to finalize first
        finally:
            # Always finalize output (even on errors) for LOCAL mode
            if mode == 'local':
                # Calculate metrics from usage metadata and timing data
                metrics = {}
                if usage_metadata_list and timing_list:
                    try:
                        metrics = calculate_metrics(usage_metadata_list, timing_list)
                        logging.info(f"Calculated metrics for finalization: {metrics}")
                    except Exception as metrics_error:
                        logging.warning(f"Failed to calculate metrics: {metrics_error}")
                        import traceback
                        logging.warning(traceback.format_exc())
                        metrics = {}
                else:
                    logging.warning(f"Cannot calculate metrics: usage_metadata_list={len(usage_metadata_list) if usage_metadata_list else 0} items, timing_list={len(timing_list) if timing_list else 0} items")
                
                try:
                    output.finalize(transcribed_pages, metrics, start_time, end_time, error_info)
                    logging.info(f"Output finalized for LOCAL mode (with {'error' if error_info else 'success'})")
                except Exception as finalize_error:
                    logging.error(f"Error during output finalization: {finalize_error}")
                    logging.error(f"Finalize traceback:\n{traceback.format_exc()}")
                
                # Log comprehensive Run Summary for LOCAL mode
                _log_run_summary(output, transcribed_pages, metrics, start_time, end_time, error_info, mode)
        
        # Re-raise the original exception if there was one (outside finally block)
        if caught_exception:
            raise caught_exception
    
    except Exception as e:
        # Final catch-all for any unhandled exceptions
        # This will catch the re-raised exception from above or any other unexpected errors
        logging.error(f"Fatal error: {str(e)}")
        logging.error(f"Full traceback:\n{traceback.format_exc()}")
        raise


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Genealogical Record Transcription Tool',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Example usage:
  python transcribe.py                    # Run interactive wizard (default)
  python transcribe.py config/config.yaml  # Use config file
  python transcribe.py --wizard-off config/config.yaml  # Explicitly disable wizard
  
Wizard mode is enabled by default. The wizard guides you through creating a configuration.
Use --wizard-off to disable wizard mode and require a config file.

The config file must be a YAML file containing all required configuration parameters.
See config/config.yaml.example for a template.
        """
    )
    parser.add_argument(
        'config_file',
        type=str,
        nargs='?',  # Optional - wizard mode is default
        help='Path to YAML configuration file (e.g., config/config.yaml). If not provided, wizard mode runs by default.'
    )
    parser.add_argument(
        '--wizard-off',
        action='store_true',
        help='Disable wizard mode and require a config file (traditional mode)'
    )
    parser.add_argument(
        '--output',
        type=str,
        help='Output path for generated config file (wizard mode only)'
    )
    
    args = parser.parse_args()
    
    # Determine if wizard mode should run
    # Wizard mode runs by default unless:
    # 1. --wizard-off is explicitly set, OR
    # 2. A config file is provided (traditional mode)
    run_wizard = not args.wizard_off and not args.config_file
    
    # Handle wizard mode (default behavior)
    if run_wizard:
        try:
            from wizard.wizard_controller import WizardController
            from wizard.steps.mode_selection_step import ModeSelectionStep
            from wizard.steps.processing_settings_step import ProcessingSettingsStep
            
            # Create wizard controller
            controller = WizardController()
            
            # Add steps
            controller.add_step(ModeSelectionStep(controller))
            from wizard.steps.context_collection_step import ContextCollectionStep
            controller.add_step(ContextCollectionStep(controller))
            controller.add_step(ProcessingSettingsStep(controller))
            
            # Run wizard
            config_path = controller.run(output_path=args.output)
            
            if config_path:
                # Continue with normal flow using generated config
                config = load_config(config_path)
                
                # Detect mode for logging purposes
                mode = detect_mode(config)
                
                # Set up logging
                log_filename, ai_log_filename, ai_logger = setup_logging(config)
                
                logging.info(f"Configuration loaded from: {config_path}")
                logging.info(f"Detected mode: {mode.upper()}")
                logging.info(f"Logging initialized. Main log: {log_filename}, AI log: {ai_log_filename}")
                
                # Load prompt text (supports both template assembly and legacy file loading)
                prompt_text = load_prompt_text(config)
                prompt_source = config.get('prompt_template') or config.get('prompt_file', 'unknown')
                logging.info(f"Prompt loaded: {prompt_source} ({len(prompt_text)} characters)")
                
                # Get logs directory
                logs_dir = "logs"
                
                # Run main function
                main(config, prompt_text, ai_logger, logs_dir, log_filename, ai_log_filename)
            else:
                print("Wizard cancelled or failed. Exiting.")
                sys.exit(1)
                
        except ImportError as e:
            print(f"Error: Wizard mode not available: {e}", file=sys.stderr)
            print("Please ensure all wizard dependencies are installed.", file=sys.stderr)
            sys.exit(1)
        except Exception as e:
            print(f"Error in wizard mode: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc()
            sys.exit(1)
    
    # Traditional mode (config file provided or wizard-off flag set)
    elif args.config_file or args.wizard_off:
        if not args.config_file:
            print("Error: --wizard-off requires a configuration file.", file=sys.stderr)
            print("Usage: python transcribe.py --wizard-off config/config.yaml", file=sys.stderr)
            sys.exit(1)
        
        try:
            # Load configuration first (needed for logging setup)
            config = load_config(args.config_file)
            
            # Detect mode for logging purposes (before logging setup to determine log filename)
            mode = detect_mode(config)
            
            # Set up logging - MUST be done before any logging calls
            log_filename, ai_log_filename, ai_logger = setup_logging(config)
            
            # Now we can log
            logging.info(f"Configuration loaded from: {args.config_file}")
            logging.info(f"Detected mode: {mode.upper()}")
            logging.info(f"Logging initialized. Main log: {log_filename}, AI log: {ai_log_filename}")
            
            # Load prompt text (supports both template assembly and legacy file loading)
            prompt_text = load_prompt_text(config)
            prompt_source = config.get('prompt_template') or config.get('prompt_file', 'unknown')
            logging.info(f"Prompt loaded: {prompt_source} ({len(prompt_text)} characters)")
            
            # Get logs directory for save_transcription_locally
            logs_dir = "logs"
            
            # Run main function
            main(config, prompt_text, ai_logger, logs_dir, log_filename, ai_log_filename)
        except FileNotFoundError as e:
            print(f"Error: {e}", file=sys.stderr)
            print(f"Please provide a valid configuration file path.", file=sys.stderr)
            sys.exit(1)
        except KeyError as e:
            print(f"Error: Missing required configuration key: {e}", file=sys.stderr)
            print(f"Please check your configuration file against config/config.yaml.example", file=sys.stderr)
            sys.exit(1)
        except yaml.YAMLError as e:
            print(f"Error: Invalid YAML in configuration file: {e}", file=sys.stderr)
            sys.exit(1)
        except Exception as e:
            print(f"Error: {e}", file=sys.stderr)
            logging.error(f"Fatal error: {e}")
            logging.error(f"Full traceback:\n{traceback.format_exc()}")
            sys.exit(1)
    else:
        # This should not happen with new default behavior, but keep as fallback
        parser.print_help()
        sys.exit(1)
